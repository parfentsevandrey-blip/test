#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор еженедельного отчёта по рынку недвижимости Нидерландов.

Читает структурированный JSON (подготовленный ИИ-агентом-исследователем) и
рендерит из него аккуратный, оформленный файл Word (.docx).

Пример запуска:
    python3 generate_report.py \
        --data data/week_2026-06-30.json \
        --out  reports/2026-06-30_dutch_realestate_RU.docx \
        --history data/history.json

Скрипт делает две вещи:
  1. Рендерит .docx из data-файла.
  2. Дописывает ключи всех новостей в history.json, чтобы в следующих
     отчётах эти материалы не повторялись (дедупликация).

Зависимости: python-docx  (pip install -r requirements.txt)
"""

import argparse
import json
import hashlib
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

try:
    import charts as _charts            # модуль графиков (matplotlib)
except Exception:                       # noqa: BLE001 — рендер без графиков, если нет matplotlib
    _charts = None

# --------------------------------------------------------------------------- #
#  Цветовая палитра (профессиональная, «финансово-аналитическая»)
# --------------------------------------------------------------------------- #
NAVY        = "1F3A5F"   # основной тёмно-синий (заголовки, обложка)
NAVY_RGB    = RGBColor(0x1F, 0x3A, 0x5F)
INK         = "222B36"   # основной цвет текста
INK_RGB     = RGBColor(0x22, 0x2B, 0x36)
MUTED       = "6B7785"   # серый для подписей/источников
MUTED_RGB   = RGBColor(0x6B, 0x77, 0x85)
LINK        = "1155CC"   # цвет ссылок
HAIRLINE    = "D7DEE6"   # тонкие линии

# Акцентные цвета сегментов
SEG_COLORS = {
    "residential": {"bar": "16846F", "soft": "E7F4F0", "text_rgb": RGBColor(0x10, 0x5A, 0x4C)},  # тёмно-бирюзовый
    "commercial":  {"bar": "C0791C", "soft": "FBF1DF", "text_rgb": RGBColor(0x8A, 0x55, 0x10)},  # янтарный
    "industrial":  {"bar": "2C5F8A", "soft": "E8F0F7", "text_rgb": RGBColor(0x1E, 0x44, 0x66)},  # стальной синий
}
DEFAULT_SEG = {"bar": NAVY, "soft": "EEF2F6", "text_rgb": NAVY_RGB}

CONCL_FILL   = "EAF3EE"   # фон блока «Вывод»
CONCL_BAR    = "16846F"   # акцентная полоса блока «Вывод»
STAT_FILL    = "F4F6F8"   # фон таблицы статистики

KT_FILL      = "1F3A5F"   # фон блока «Главные выводы» (тёмный)
WATCH_FILL   = "FFF6E5"   # фон врезки «За чем следить»
WATCH_BAR    = "C0791C"   # полоса врезки «За чем следить»
OUTLOOK_FILL = "EEF2F6"   # фон блока «Картина и прогноз»
GLOSS_FILL   = "F4F6F8"   # фон словаря терминов

# Теги влияния: (символ, фон, цвет текста, подпись)
DIRECTION = {
    "up":      ("▲", "E7F4EE", RGBColor(0x1E, 0x7A, 0x4D), "возможность / рост"),
    "down":    ("▼", "FBEAEA", RGBColor(0xB0, 0x3A, 0x3A), "риск / снижение"),
    "neutral": ("◆", "E8F0F7", RGBColor(0x2C, 0x5F, 0x8A), "структурный сдвиг"),
}

# Статусы сюжетов в развитии (память сюжетов «не повторяться → продолжать»)
THREAD_STATUS = {
    "new":        ("НОВЫЙ",            "E8F0F7", RGBColor(0x2C, 0x5F, 0x8A)),
    "developing": ("В РАЗВИТИИ",       "FBF1DF", RGBColor(0x8A, 0x55, 0x10)),
    "watch":      ("ПОД НАБЛЮДЕНИЕМ",  "FFF6E5", RGBColor(0xB0, 0x7A, 0x1C)),
    "resolved":   ("ЗАКРЫТ",          "E7F4EE", RGBColor(0x1E, 0x7A, 0x4D)),
}
KIND_ICON = {"deal_deadline": "⏳", "cbs_release": "📈", "vote": "🗳",
             "reit_earnings": "📊", "other": "•"}
PORT_FILL = "FBF1DF"   # фон врезки «Важно для вашего портфеля»
SEG_RU_SHORT = {"residential": "жильё", "commercial": "ритейл",
                "industrial": "индустриал", "overview": "все", "macro": "макро"}

BASE_FONT = "Calibri"

RU_MONTHS = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
    7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}

# Порядок и человекочитаемые названия подразделов
SUBSECTION_ORDER = ["laws", "news", "trends", "stats"]
SUBSECTION_TITLES = {
    "laws":   "Изменения в законах и регулировании",
    "news":   "Новости",
    "trends": "Тренды",
    "stats":  "Статистика",
}


# --------------------------------------------------------------------------- #
#  Низкоуровневые помощники для оформления (OOXML)
# --------------------------------------------------------------------------- #
def _set_cell_background(cell, fill_hex):
    """Заливка ячейки таблицы цветом."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    """Внутренние отступы ячейки (в twips: 20 = 1pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left),
                     ("end", right), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def _set_cell_borders(cell, edges):
    """edges: dict edge -> (sz_eighths_pt, color_hex). edge in top/bottom/left/right."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in edges:
            sz, color = edges[edge]
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
    tcPr.append(borders)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def _no_row_split(table):
    """Запретить разрыв строк таблицы между страницами (боксы не рвутся)."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantsplit = OxmlElement("w:cantSplit")
        trPr.append(cantsplit)


def _keep_with_next(paragraph):
    """Держать абзац вместе со следующим (заголовок не висит внизу страницы)."""
    paragraph.paragraph_format.keep_with_next = True


def _full_width(table, width_cm=16.8):
    """Растянуть таблицу на ширину текста (полезно для цветных полос/боксов)."""
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))  # 1 см = 567 twips
    tblPr.append(tblW)
    for row in table.rows:
        for c in row.cells:
            c.width = Cm(width_cm / len(row.cells))


def _shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _add_chip(paragraph, text, fill_hex, color_rgb, size=8):
    """Маленькая цветная «плашка» (chip) — фон + цветной текст на уровне run."""
    run = paragraph.add_run(f" {text} ")
    _style_run(run, size=size, bold=True, color_rgb=color_rgb)
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    rPr.append(shd)
    return run


def _add_hyperlink(paragraph, url, text, color=LINK, size=9, underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), BASE_FONT); rf.set(qn("w:hAnsi"), BASE_FONT); rPr.append(rf)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_page_number_field(paragraph):
    """Вставить поле «Стр. X из Y» в колонтитул."""
    def field(instr):
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve"); instr_el.text = instr
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        return fld_begin, instr_el, fld_end

    run = paragraph.add_run("Стр. ")
    run.font.size = Pt(8); run.font.color.rgb = MUTED_RGB; run.font.name = BASE_FONT
    for el in field("PAGE"):
        run._r.append(el)
    run2 = paragraph.add_run(" из ")
    run2.font.size = Pt(8); run2.font.color.rgb = MUTED_RGB; run2.font.name = BASE_FONT
    for el in field("NUMPAGES"):
        run2._r.append(el)


# --------------------------------------------------------------------------- #
#  Помощники-конструкторы блоков отчёта
# --------------------------------------------------------------------------- #
def _style_run(run, size=10.5, bold=False, italic=False, color_rgb=INK_RGB, font=BASE_FONT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color_rgb
    run.font.name = font
    # Для корректного отображения кириллицы в некоторых вьюверах
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font)
    return run


def _p(doc, space_before=0, space_after=4, line=1.12, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pf.alignment = align
    return p


def add_segment_bar(doc, title, color_hex, icon=""):
    """Цветная полоса-заголовок сегмента на всю ширину."""
    _p(doc, space_before=10, space_after=0)
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, color_hex)
    _set_cell_margins(cell, top=120, bottom=120, left=200, right=200)
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label = f"{icon}  {title}" if icon else title
    run = p.add_run(label)
    _style_run(run, size=14, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    _keep_with_next(p)
    _no_row_split(table)


def add_subsection_title(doc, text, color_rgb):
    p = _p(doc, space_before=10, space_after=3)
    run = p.add_run(text)
    _style_run(run, size=11.5, bold=True, color_rgb=color_rgb)
    _keep_with_next(p)
    # тонкая нижняя линия
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), HAIRLINE)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_bullet(doc, text, source=None, url=None, date=None, impact=None, direction=None):
    """Пункт-маркер: факт + (опц.) врезка «Почему важно» с тегом влияния + источник."""
    sym, fill, drgb, _label = DIRECTION.get(direction or "", ("▪", None, MUTED_RGB, ""))
    p = _p(doc, space_before=3, space_after=1, line=1.12)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    bullet = p.add_run(f"{sym}  ")
    _style_run(bullet, size=10.5, bold=True, color_rgb=drgb)
    body = p.add_run(text)
    _style_run(body, size=10.5, color_rgb=INK_RGB)

    # врезка «Почему важно» с цветом тега влияния
    if impact:
        ip = _p(doc, space_before=1, space_after=1, line=1.1)
        ip.paragraph_format.left_indent = Cm(0.5)
        if fill:
            _shade_paragraph(ip, fill)
        lab = ip.add_run("  Почему важно:  ")
        _style_run(lab, size=9, bold=True, color_rgb=drgb)
        r = ip.add_run(impact + "  ")
        _style_run(r, size=9, color_rgb=INK_RGB)

    # подпись-источник под пунктом
    if source or url or date:
        meta = _p(doc, space_before=0, space_after=5, line=1.0)
        meta.paragraph_format.left_indent = Cm(0.5)
        parts = []
        if date:
            parts.append(f"📅 {date}")
        if source:
            parts.append(f"Источник: {source}")
        prefix = "   " + "   •   ".join(parts) if parts else "   "
        r = meta.add_run(prefix)
        _style_run(r, size=8.5, italic=True, color_rgb=MUTED_RGB)
        if url:
            sep = meta.add_run("   ")
            _style_run(sep, size=8.5, color_rgb=MUTED_RGB)
            _add_hyperlink(meta, url, "ссылка ↗", size=8.5)


def add_stats_table(doc, items):
    """Подраздел «Статистика» в виде таблицы: показатель | значение | источник."""
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _full_width(table)
    hdr = table.rows[0].cells
    for i, htext in enumerate(("Показатель", "Значение / динамика")):
        _set_cell_background(hdr[i], NAVY)
        _set_cell_margins(hdr[i])
        r = hdr[i].paragraphs[0].add_run(htext)
        _style_run(r, size=9.5, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    # ширины колонок
    table.columns[0].width = Cm(10.0)
    table.columns[1].width = Cm(6.8)

    for idx, it in enumerate(items):
        row = table.add_row().cells
        fill = STAT_FILL if idx % 2 == 0 else "FFFFFF"
        metric = it.get("text") or it.get("metric") or ""
        value = it.get("value", "")
        sym, _f, drgb, _l = DIRECTION.get(it.get("direction") or "", ("", None, INK_RGB, ""))
        for ci in (0, 1):
            _set_cell_background(row[ci], fill)
            _set_cell_margins(row[ci])
            row[ci].paragraphs[0].paragraph_format.space_after = Pt(0)
        # колонка 1: показатель + (опц.) «почему важно»
        r = row[0].paragraphs[0].add_run(metric)
        _style_run(r, size=9.5, color_rgb=INK_RGB)
        if it.get("impact"):
            ip = row[0].add_paragraph()
            ip.paragraph_format.space_before = Pt(1); ip.paragraph_format.space_after = Pt(0)
            r = ip.add_run("→ " + it["impact"])
            _style_run(r, size=8, italic=True, color_rgb=MUTED_RGB)
        # колонка 2: стрелка тренда + значение
        vp = row[1].paragraphs[0]
        if sym:
            ar = vp.add_run(sym + " ")
            _style_run(ar, size=9.5, bold=True, color_rgb=drgb)
        r = vp.add_run(value)
        _style_run(r, size=9.5, bold=True, color_rgb=INK_RGB)
        if it.get("source") or it.get("url"):
            mp = row[1].add_paragraph()
            mp.paragraph_format.space_before = Pt(1); mp.paragraph_format.space_after = Pt(0)
            if it.get("source"):
                r = mp.add_run(it["source"])
                _style_run(r, size=7.5, italic=True, color_rgb=MUTED_RGB)
            if it.get("url"):
                _add_hyperlink(mp, it["url"], "  ↗", size=7.5)
    _p(doc, space_after=2)


def add_conclusion_box(doc, text):
    """Блок «Вывод» — заметный бокс с акцентной левой полосой."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, CONCL_FILL)
    _set_cell_margins(cell, top=140, bottom=140, left=220, right=220)
    _set_cell_borders(cell, {"left": (36, CONCL_BAR)})  # толстая левая полоса
    _no_row_split(table)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    label = p0.add_run("ВЫВОД")
    _style_run(label, size=9, bold=True, color_rgb=RGBColor(0x10, 0x5A, 0x4C))

    p1 = cell.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = 1.12
    r = p1.add_run(text)
    _style_run(r, size=10.5, color_rgb=INK_RGB)
    _p(doc, space_after=6)


def add_watch_box(doc, text):
    """Врезка «За чем следить» — компактный бокс с янтарной левой полосой."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, WATCH_FILL)
    _set_cell_margins(cell, top=100, bottom=100, left=220, right=220)
    _set_cell_borders(cell, {"left": (30, WATCH_BAR)})
    _no_row_split(table)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    p0.paragraph_format.line_spacing = 1.1
    lab = p0.add_run("👁  ЗА ЧЕМ СЛЕДИТЬ.  ")
    _style_run(lab, size=9, bold=True, color_rgb=RGBColor(0x8A, 0x55, 0x10))
    r = p0.add_run(text)
    _style_run(r, size=10, color_rgb=INK_RGB)
    _p(doc, space_after=8)


def build_key_takeaways(doc, data):
    """Тёмный блок «Главные выводы недели» — нумерованные тезисы."""
    kt = data.get("key_takeaways") or []
    if not kt:
        return
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, KT_FILL)
    _no_row_split(table)
    _set_cell_margins(cell, top=150, bottom=150, left=220, right=220)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(4)
    hr = p0.add_run("ГЛАВНЫЕ ВЫВОДЫ НЕДЕЛИ")
    _style_run(hr, size=11, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    for i, t in enumerate(kt, 1):
        pp = cell.add_paragraph()
        pp.paragraph_format.space_before = Pt(2)
        pp.paragraph_format.space_after = Pt(2)
        pp.paragraph_format.line_spacing = 1.15
        pp.paragraph_format.left_indent = Cm(0.55)
        pp.paragraph_format.first_line_indent = Cm(-0.55)
        num = pp.add_run(f"{i}.  ")
        _style_run(num, size=10.5, bold=True, color_rgb=RGBColor(0xF0, 0xC4, 0x6A))
        r = pp.add_run(t)
        _style_run(r, size=10.5, color_rgb=RGBColor(0xEC, 0xF1, 0xF6))
    _p(doc, space_after=8)


def build_outlook(doc, data):
    """Блок «Картина недели и прогноз»."""
    outlook = data.get("outlook")
    if not outlook:
        return
    add_segment_bar(doc, "Картина недели и прогноз", "1F3A5F", "🧭")
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, OUTLOOK_FILL)
    _no_row_split(table)
    _set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.line_spacing = 1.18
    r = cp.add_run(outlook)
    _style_run(r, size=10.5, color_rgb=INK_RGB)
    _p(doc, space_after=6)


def build_glossary(doc, data):
    """Словарь терминов — двухколоночная таблица термин/определение."""
    gl = data.get("glossary") or []
    if not gl:
        return
    add_segment_bar(doc, "Словарь терминов", "5D6D7E", "📖")
    table = doc.add_table(rows=0, cols=2)
    _remove_table_borders(table)
    _full_width(table)
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(11.8)
    for idx, g in enumerate(gl):
        row = table.add_row().cells
        fill = GLOSS_FILL if idx % 2 == 0 else "FFFFFF"
        for ci in (0, 1):
            _set_cell_background(row[ci], fill)
            _set_cell_margins(row[ci], top=60, bottom=60)
            row[ci].paragraphs[0].paragraph_format.space_after = Pt(0)
        tr = row[0].paragraphs[0].add_run(g.get("term", ""))
        _style_run(tr, size=9.5, bold=True, color_rgb=NAVY_RGB)
        dr = row[1].paragraphs[0].add_run(g.get("definition", ""))
        _style_run(dr, size=9.5, color_rgb=INK_RGB)
    _p(doc, space_after=6)


# --------------------------------------------------------------------------- #
#  Обложка и служебные секции
# --------------------------------------------------------------------------- #
def fmt_period(week_start, week_end):
    try:
        d1 = datetime.strptime(week_start, "%Y-%m-%d")
        d2 = datetime.strptime(week_end, "%Y-%m-%d")
        if d1.month == d2.month:
            return f"{d1.day}–{d2.day} {RU_MONTHS[d2.month]} {d2.year} г."
        return f"{d1.day} {RU_MONTHS[d1.month]} – {d2.day} {RU_MONTHS[d2.month]} {d2.year} г."
    except Exception:
        return f"{week_start} — {week_end}"


def build_cover(doc, data):
    period = fmt_period(data.get("week_start", ""), data.get("week_end", ""))

    # верхний отступ
    for _ in range(2):
        _p(doc, space_after=0)

    # надзаголовок
    p = _p(doc, space_after=2)
    r = p.add_run("ЕЖЕНЕДЕЛЬНАЯ АНАЛИТИКА")
    _style_run(r, size=11, bold=True, color_rgb=MUTED_RGB)
    # межбуквенный интервал
    rPr = r._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "60"); rPr.append(spc)

    # основной заголовок
    p = _p(doc, space_after=2)
    r = p.add_run("Рынок недвижимости")
    _style_run(r, size=30, bold=True, color_rgb=NAVY_RGB)
    p = _p(doc, space_after=10)
    r = p.add_run("Нидерландов")
    _style_run(r, size=30, bold=True, color_rgb=NAVY_RGB)

    # цветная полоса-разделитель
    table = doc.add_table(rows=1, cols=3)
    _remove_table_borders(table)
    _full_width(table)
    for i, key in enumerate(["residential", "commercial", "industrial"]):
        _set_cell_background(table.cell(0, i), SEG_COLORS[key]["bar"])
        table.cell(0, i).paragraphs[0].add_run(" ").font.size = Pt(3)
        table.cell(0, i).paragraphs[0].paragraph_format.space_after = Pt(0)
    _p(doc, space_after=8)

    # заголовок недели (headline)
    headline = data.get("headline")
    if headline:
        hp = _p(doc, space_before=2, space_after=10)
        r = hp.add_run(f"«{headline}»")
        _style_run(r, size=14, bold=True, italic=True, color_rgb=SEG_COLORS["residential"]["text_rgb"])

    # период
    p = _p(doc, space_after=2)
    r = p.add_run("Период: ")
    _style_run(r, size=13, bold=True, color_rgb=INK_RGB)
    r = p.add_run(period)
    _style_run(r, size=13, color_rgb=INK_RGB)

    # сегменты
    p = _p(doc, space_after=2)
    r = p.add_run("Сегменты: ")
    _style_run(r, size=11, bold=True, color_rgb=MUTED_RGB)
    r = p.add_run("жилая недвижимость · коммерция (стрит-ритейл) · индустриальная (склады, промзоны)")
    _style_run(r, size=11, color_rgb=MUTED_RGB)

    # дата формирования
    gen = data.get("report_date") or datetime.now().strftime("%Y-%m-%d")
    try:
        gd = datetime.strptime(gen, "%Y-%m-%d")
        gen_h = f"{gd.day} {RU_MONTHS[gd.month]} {gd.year} г."
    except Exception:
        gen_h = gen
    p = _p(doc, space_before=2, space_after=14)
    r = p.add_run(f"Отчёт сформирован: {gen_h}")
    _style_run(r, size=9.5, italic=True, color_rgb=MUTED_RGB)


def build_exec_summary(doc, data):
    summary = data.get("executive_summary")
    if not summary:
        return
    add_subsection_title(doc, "Коротко о неделе", NAVY_RGB)
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, "EEF2F6")
    _set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.line_spacing = 1.15
    r = cp.add_run(summary)
    _style_run(r, size=10.5, color_rgb=INK_RGB)
    _p(doc, space_after=6)


def build_segment(doc, seg, charts_list=None, pngs=None):
    colors = SEG_COLORS.get(seg.get("id"), DEFAULT_SEG)
    add_segment_bar(doc, seg.get("title", ""), colors["bar"], seg.get("icon", ""))

    subs = seg.get("subsections", {})
    # subsections может быть dict (по ключам) или list
    def get_items(key):
        if isinstance(subs, dict):
            return subs.get(key) or []
        for s in subs:
            if s.get("id") == key or s.get("key") == key:
                return s.get("items") or []
        return []

    any_content = False
    for key in SUBSECTION_ORDER:
        items = get_items(key)
        if not items:
            continue
        any_content = True
        add_subsection_title(doc, SUBSECTION_TITLES[key], colors["text_rgb"])
        if key == "stats":
            add_stats_table(doc, items)
        else:
            for it in items:
                add_bullet(
                    doc,
                    it.get("text", ""),
                    source=it.get("source"),
                    url=it.get("url"),
                    date=it.get("date"),
                    impact=it.get("impact"),
                    direction=it.get("direction"),
                )

    if not any_content:
        p = _p(doc, space_after=4)
        r = p.add_run("За отчётную неделю значимых событий по этому сегменту не зафиксировано.")
        _style_run(r, size=10.5, italic=True, color_rgb=MUTED_RGB)

    # графики этого сегмента
    seg_charts = [c for c in (charts_list or []) if c.get("segment") == seg.get("id")]
    if seg_charts:
        build_charts(doc, seg_charts, pngs)

    concl = seg.get("conclusion")
    if concl:
        add_conclusion_box(doc, concl)

    watch = seg.get("watch")
    if watch:
        add_watch_box(doc, watch)


def add_chart_image(doc, png_path, caption=None):
    """Встроить PNG-график по ширине текста + подпись-вывод под ним."""
    doc.add_picture(png_path, width=Cm(16.4))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    if caption:
        _keep_with_next(p)  # картинка держится со своей подписью
    if caption:
        cp = _p(doc, space_before=0, space_after=8, line=1.08, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = cp.add_run(caption)
        _style_run(r, size=9, italic=True, color_rgb=MUTED_RGB)


def build_charts(doc, specs, pngs):
    """Отрисовать список графиков (по готовым PNG) с подписями."""
    for spec in specs or []:
        path = (pngs or {}).get(spec.get("id"))
        if path:
            add_chart_image(doc, path, spec.get("caption"))


def build_overview_charts(doc, data, pngs):
    ov = [c for c in (data.get("charts") or []) if c.get("segment") == "overview"]
    if not ov:
        return
    add_segment_bar(doc, "Статистика в графиках", "1F3A5F", "📊")
    _p(doc, space_after=2)
    build_charts(doc, ov, pngs)


def build_threads(doc, data):
    """«Сюжеты в развитии» — память историй со статус-бейджами и след. триггером."""
    threads = data.get("threads") or []
    if not threads:
        return
    add_segment_bar(doc, "Сюжеты в развитии", "5D6D7E", "🧵")
    _p(doc, space_after=2)
    for t in threads:
        st = THREAD_STATUS.get((t.get("status") or "").lower(), ("—", "EEF2F6", MUTED_RGB))
        p = _p(doc, space_before=5, space_after=1, line=1.12)
        _add_chip(p, st[0], st[1], st[2], size=8)
        r = p.add_run("  " + t.get("title", ""))
        _style_run(r, size=11, bold=True, color_rgb=NAVY_RGB)
        _keep_with_next(p)
        if t.get("update"):
            up = _p(doc, space_before=0, space_after=1, line=1.1)
            up.paragraph_format.left_indent = Cm(0.2)
            r = up.add_run("Что нового: ")
            _style_run(r, size=9.5, bold=True, color_rgb=MUTED_RGB)
            r = up.add_run(t["update"])
            _style_run(r, size=9.5, color_rgb=INK_RGB)
        nt = t.get("next_trigger") or {}
        if nt.get("date") or nt.get("what"):
            tp = _p(doc, space_before=0, space_after=4, line=1.05)
            tp.paragraph_format.left_indent = Cm(0.2)
            r = tp.add_run("Следующий триггер: ")
            _style_run(r, size=9, bold=True, color_rgb=RGBColor(0x8A, 0x55, 0x10))
            trig = f"{nt.get('date', '')} — {nt.get('what', '')}".strip(" —")
            r = tp.add_run(trig)
            _style_run(r, size=9, italic=True, color_rgb=MUTED_RGB)
            if t.get("url"):
                _add_hyperlink(tp, t["url"], "  ↗", size=8.5)
    _p(doc, space_after=4)


def build_portfolio(doc, data):
    """Врезка «Важно для вашего портфеля» — персонализация по profile.json."""
    notes = data.get("portfolio_notes") or []
    if not notes:
        return
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _full_width(table)
    cell = table.cell(0, 0)
    _set_cell_background(cell, PORT_FILL)
    _no_row_split(table)
    _set_cell_margins(cell, top=130, bottom=130, left=220, right=220)
    _set_cell_borders(cell, {"left": (30, "C0791C")})
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    lab = p0.add_run("★  ВАЖНО ДЛЯ ВАШЕГО ПОРТФЕЛЯ")
    _style_run(lab, size=10, bold=True, color_rgb=RGBColor(0x8A, 0x55, 0x10))
    for n in notes:
        text = n.get("text", "") if isinstance(n, dict) else str(n)
        pp = cell.add_paragraph()
        pp.paragraph_format.space_before = Pt(2)
        pp.paragraph_format.space_after = Pt(2)
        pp.paragraph_format.line_spacing = 1.15
        pp.paragraph_format.left_indent = Cm(0.5)
        pp.paragraph_format.first_line_indent = Cm(-0.35)
        b = pp.add_run("▪  ")
        _style_run(b, size=10.5, bold=True, color_rgb=RGBColor(0x8A, 0x55, 0x10))
        r = pp.add_run(text)
        _style_run(r, size=10.5, color_rgb=INK_RGB)
    _p(doc, space_after=6)


def build_calendar(doc, data):
    """«Календарь: за чем следить» — форвард-вотчлист ближайших триггеров."""
    cal = data.get("calendar") or []
    if not cal:
        return

    def _key(c):
        try:
            return (0, datetime.strptime(c.get("date", ""), "%Y-%m-%d"))
        except Exception:
            return (1, datetime.max)
    cal = sorted(cal, key=_key)

    add_segment_bar(doc, "Календарь: за чем следить", "C0791C", "📅")
    _p(doc, space_after=2)
    table = doc.add_table(rows=1, cols=3)
    _remove_table_borders(table)
    _full_width(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(("Дата", "Событие", "Сегмент")):
        _set_cell_background(hdr[i], NAVY)
        _set_cell_margins(hdr[i])
        r = hdr[i].paragraphs[0].add_run(h)
        _style_run(r, size=9.5, bold=True, color_rgb=RGBColor(0xFF, 0xFF, 0xFF))
    table.columns[0].width = Cm(2.7)
    table.columns[1].width = Cm(11.3)
    table.columns[2].width = Cm(2.8)
    for idx, c in enumerate(cal):
        row = table.add_row().cells
        fill = STAT_FILL if idx % 2 == 0 else "FFFFFF"
        icon = KIND_ICON.get(c.get("kind", "other"), "•")
        cells = (c.get("date", ""), f"{icon}  {c.get('what', '')}", SEG_RU_SHORT.get(c.get("segment"), c.get("segment", "")))
        for ci, txt in enumerate(cells):
            _set_cell_background(row[ci], fill)
            _set_cell_margins(row[ci])
            cp = row[ci].paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            r = cp.add_run(txt)
            _style_run(r, size=9.5, bold=(ci == 0), color_rgb=INK_RGB)
        if c.get("impact"):
            mp = row[1].add_paragraph()
            mp.paragraph_format.space_before = Pt(1); mp.paragraph_format.space_after = Pt(0)
            r = mp.add_run("→ " + c["impact"])
            _style_run(r, size=8, italic=True, color_rgb=MUTED_RGB)
    _p(doc, space_after=6)


def build_sources(doc, data):
    sources = data.get("sources") or []
    if not sources:
        return
    add_segment_bar(doc, "Источники", "5D6D7E", "🔗")
    _p(doc, space_after=2)
    for s in sources:
        p = _p(doc, space_before=1, space_after=2, line=1.05)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        b = p.add_run("•  ")
        _style_run(b, size=9.5, color_rgb=MUTED_RGB)
        name = s.get("name", s.get("url", ""))
        r = p.add_run(f"{name}  ")
        _style_run(r, size=9.5, color_rgb=INK_RGB)
        if s.get("url"):
            _add_hyperlink(p, s["url"], s["url"], size=9)


def build_footer_disclaimer(doc):
    _p(doc, space_before=10, space_after=0)
    p = _p(doc, space_before=6, space_after=0)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4"); top.set(qn("w:color"), HAIRLINE)
    pbdr.append(top); pPr.append(pbdr)
    r = p.add_run(
        "Дисклеймер: материал носит информационно-аналитический характер, "
        "подготовлен на основе открытых источников и не является инвестиционной "
        "рекомендацией. Проверяйте данные по первоисточникам перед принятием решений."
    )
    _style_run(r, size=8, italic=True, color_rgb=MUTED_RGB)


# --------------------------------------------------------------------------- #
#  Глобальные стили и колонтитулы
# --------------------------------------------------------------------------- #
def setup_document(data):
    doc = Document()
    # базовый стиль
    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK_RGB
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:cs"), BASE_FONT)

    # поля страницы
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # верхний колонтитул
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    period = fmt_period(data.get("week_start", ""), data.get("week_end", ""))
    r = hp.add_run(f"Недвижимость Нидерландов · {period}")
    _style_run(r, size=8, color_rgb=MUTED_RGB)

    # нижний колонтитул со страницами
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(fp)
    return doc


# --------------------------------------------------------------------------- #
#  Дедупликация
# --------------------------------------------------------------------------- #
def norm_key(item):
    """Стабильный ключ новости для дедупа: по URL, иначе по нормализованному тексту."""
    url = (item.get("url") or "").strip().lower()
    url = re.sub(r"[#?].*$", "", url).rstrip("/")
    if url:
        return "u:" + url
    text = (item.get("text") or item.get("metric") or "").strip().lower()
    text = re.sub(r"\s+", " ", text)
    return "t:" + hashlib.sha1(text.encode("utf-8")).hexdigest()[:16]


def iter_items(data):
    for seg in data.get("segments", []):
        subs = seg.get("subsections", {})
        if isinstance(subs, dict):
            for key, items in subs.items():
                for it in (items or []):
                    yield seg.get("id"), key, it
        else:
            for s in subs:
                for it in (s.get("items") or []):
                    yield seg.get("id"), s.get("id"), it


def load_history(path):
    if not path or not os.path.exists(path):
        return {"items": {}, "weeks": []}
    try:
        with open(path, "r", encoding="utf-8") as f:
            h = json.load(f)
        h.setdefault("items", {})
        h.setdefault("weeks", [])
        return h
    except Exception:
        return {"items": {}, "weeks": []}


def update_history(path, data):
    if not path:
        return 0, 0
    h = load_history(path)
    added = 0
    total = 0
    report_date = data.get("report_date") or datetime.now().strftime("%Y-%m-%d")
    for seg_id, sub, it in iter_items(data):
        total += 1
        k = norm_key(it)
        if k not in h["items"]:
            h["items"][k] = {
                "date": report_date,
                "segment": seg_id,
                "subsection": sub,
                "title": (it.get("text") or it.get("metric") or "")[:160],
                "url": it.get("url", ""),
            }
            added += 1
    h["weeks"].append({
        "report_date": report_date,
        "week_start": data.get("week_start"),
        "week_end": data.get("week_end"),
        "items": total,
        "new_items": added,
    })
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(h, f, ensure_ascii=False, indent=2)
    return added, total


def check_duplicates(history_path, data):
    """Предупредить, если в новых данных есть материалы, уже бывшие в прошлых отчётах."""
    h = load_history(history_path)
    seen = h.get("items", {})
    dups = []
    for seg_id, sub, it in iter_items(data):
        k = norm_key(it)
        if k in seen:
            dups.append((seg_id, it.get("text", "")[:80], seen[k].get("date")))
    return dups


# --------------------------------------------------------------------------- #
#  Сборка
# --------------------------------------------------------------------------- #
def build_report(data, out_path):
    # графики недели + авто-тренды из памяти (trend_chart_specs)
    charts_list = (data.get("charts") or []) + (data.get("trend_chart_specs") or [])
    data["charts"] = charts_list  # чтобы overview/segment-фильтры увидели и тренды
    pngs = {}
    if charts_list and _charts is not None:
        assets_dir = os.path.join(os.path.dirname(out_path) or ".", "assets")
        pngs = _charts.render_charts(charts_list, assets_dir)
    elif charts_list and _charts is None:
        print("⚠️  matplotlib не установлен — графики пропущены (pip install -r requirements.txt).")

    doc = setup_document(data)
    build_cover(doc, data)
    doc.add_page_break()           # обложка — отдельная страница
    build_key_takeaways(doc, data)
    build_portfolio(doc, data)     # «Важно для вашего портфеля» (персонализация)
    build_threads(doc, data)       # «Сюжеты в развитии» (память историй)
    build_overview_charts(doc, data, pngs)
    build_exec_summary(doc, data)
    for seg in data.get("segments", []):
        build_segment(doc, seg, charts_list, pngs)
    build_outlook(doc, data)
    build_calendar(doc, data)      # «Календарь: за чем следить» (форсайт)
    build_glossary(doc, data)
    build_sources(doc, data)
    build_footer_disclaimer(doc)
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="Генератор Word-отчёта по рынку недвижимости Нидерландов")
    ap.add_argument("--data", required=True, help="JSON с данными отчёта")
    ap.add_argument("--out", required=True, help="Путь к выходному .docx")
    ap.add_argument("--history", default=None, help="Путь к history.json (дедуп)")
    ap.add_argument("--no-history-update", action="store_true",
                    help="Не дописывать историю (только проверка дубликатов)")
    args = ap.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    if args.history:
        dups = check_duplicates(args.history, data)
        if dups:
            print(f"⚠️  Обнаружено {len(dups)} материалов, уже встречавшихся в прошлых отчётах:")
            for seg, title, when in dups[:15]:
                print(f"    - [{seg}] {title}…  (был в отчёте {when})")
            print("    → агенту-исследователю стоит заменить их свежими материалами.")

    out = build_report(data, args.out)
    print(f"✅ Отчёт сохранён: {out}")

    if args.history and not args.no_history_update:
        added, total = update_history(args.history, data)
        print(f"🗂  История обновлена: +{added} новых из {total} материалов "
              f"(файл: {args.history}).")


if __name__ == "__main__":
    main()
