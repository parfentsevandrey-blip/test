"""Рендер отчёта в .docx в оформлении эталонного шаблона."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

from . import style as S

log = logging.getLogger(__name__)


# --------------------------------------------------------------------------
# низкоуровневые помощники OOXML
# --------------------------------------------------------------------------
def _el(tag: str, **attrs) -> OxmlElement:
    node = OxmlElement(tag if ":" in tag else f"w:{tag}")
    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), str(value))
    return node


# Порядок дочерних элементов в OOXML жёстко задан схемой: если вставить узел
# «не туда», Word откроет файл, а LibreOffice откажется его читать.
_ORDER = {
    "pPr": (
        "pStyle keepNext keepLines pageBreakBefore framePr widowControl numPr "
        "suppressLineNumbers pBdr shd tabs suppressAutoHyphens kinsoku wordWrap "
        "overflowPunct topLinePunct autoSpaceDE autoSpaceDN bidi adjustRightInd "
        "snapToGrid spacing ind contextualSpacing mirrorIndents suppressOverlap jc "
        "textDirection textAlignment textboxTightWrap outlineLvl divId cnfStyle rPr "
        "sectPr pPrChange"
    ),
    "rPr": (
        "rStyle rFonts b bCs i iCs caps smallCaps strike dstrike outline shadow emboss "
        "imprint noProof snapToGrid vanish webHidden color spacing w kern position sz "
        "szCs highlight u effect bdr shd fitText vertAlign rtl cs em lang "
        "eastAsianLayout specVanish oMath"
    ),
    "tblPr": (
        "tblStyle tblpPr tblOverlap bidiVisual tblStyleRowBandSize tblStyleColBandSize "
        "tblW jc tblCellSpacing tblInd tblBorders shd tblLayout tblCellMar tblLook "
        "tblCaption tblDescription"
    ),
    "tcPr": (
        "cnfStyle tcW gridSpan hMerge vMerge tcBorders shd noWrap tcMar textDirection "
        "tcFitText vAlign hideMark"
    ),
}


def _insert_ordered(parent, element) -> None:
    """Вставляет узел в родителя с соблюдением порядка, заданного схемой OOXML."""
    order = _ORDER[parent.tag.split("}")[-1]].split()
    tag = element.tag.split("}")[-1]
    index = order.index(tag)
    for child in parent:
        child_tag = child.tag.split("}")[-1]
        if child_tag in order and order.index(child_tag) > index:
            child.addprevious(element)
            return
    parent.append(element)


def _get_or_add(parent, tag: str):
    existing = parent.find(qn(f"w:{tag}"))
    if existing is not None:
        return existing
    element = _el(tag)
    _insert_ordered(parent, element)
    return element


def paragraph_border(paragraph, edge: str, color: str, size: int, space: int = 1) -> None:
    """Добавляет линейку (границу абзаца) — используется для линеек и подчёркиваний."""
    borders = _get_or_add(paragraph._p.get_or_add_pPr(), "pBdr")
    borders.append(_el(edge, val="single", sz=size, space=space, color=color))


def run_tracking(run, twentieths: int) -> None:
    """Разрядка символов (w:spacing) в 1/20 pt."""
    _insert_ordered(run._r.get_or_add_rPr(), _el("spacing", val=twentieths))


def cell_shading(cell, color: str) -> None:
    _insert_ordered(
        cell._tc.get_or_add_tcPr(), _el("shd", val="clear", color="auto", fill=color)
    )


def cell_margins(cell, top=0, bottom=0, left=0, right=0) -> None:
    """Внутренние поля ячейки в pt."""
    margins = _el("tcMar")
    for name, value in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        margins.append(_el(name, w=int(value * 20), type="dxa"))
    _insert_ordered(cell._tc.get_or_add_tcPr(), margins)


def table_borders(table, spec: dict[str, tuple[str, int]]) -> None:
    """spec: {'top': (color, sz), 'insideH': (...), ...}; отсутствующие рёбра — none."""
    borders = _el("tblBorders")
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in spec:
            color, size = spec[edge]
            borders.append(_el(edge, val="single", sz=size, space=0, color=color))
        else:
            borders.append(_el(edge, val="none", sz=0, space=0, color="auto"))
    _insert_ordered(table._tbl.tblPr, borders)


def keep_with_next(paragraph) -> None:
    _insert_ordered(paragraph._p.get_or_add_pPr(), _el("keepNext", val="1"))


# --------------------------------------------------------------------------
# строительные блоки страницы
# --------------------------------------------------------------------------
def add_run(paragraph, text: str, *, size: float, color: str, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = S.FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    # шрифт для кириллицы/восточных наборов
    fonts = _get_or_add(run._r.get_or_add_rPr(), "rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), S.FONT)
    return run


def add_paragraph(doc, *, before=0, after=0, align=None, line=None):
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    if line is not None:
        fmt.line_spacing = line
    return paragraph


def gold_rule(doc, *, before=0, after=6, size=S.RULE_GOLD_SZ):
    paragraph = add_paragraph(doc, before=before, after=after)
    paragraph.paragraph_format.line_spacing = 1
    paragraph_border(paragraph, "bottom", S.GOLD, size, space=1)
    return paragraph


def section_heading(doc, text: str):
    paragraph = add_paragraph(doc, before=16, after=8)
    run = add_run(paragraph, text.upper(), size=S.FS_SECTION, color=S.NAVY, bold=True)
    run_tracking(run, S.TRACKING_SECTION)
    paragraph_border(paragraph, "bottom", S.NAVY, S.RULE_NAVY_SZ, space=4)
    keep_with_next(paragraph)
    return paragraph


def body_paragraph(doc, text: str, *, justify=True, after=6):
    align = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    paragraph = add_paragraph(doc, after=after, align=align)
    add_run(paragraph, text, size=S.FS_BODY, color=S.BODY)
    return paragraph


def bullet(doc, text: str):
    paragraph = add_paragraph(doc, after=3)
    fmt = paragraph.paragraph_format
    fmt.left_indent = Pt(S.BULLET_INDENT)
    fmt.first_line_indent = Pt(-S.BULLET_HANGING)
    tabs = _el("tabs")
    tabs.append(_el("tab", val="left", pos=int(S.BULLET_INDENT * 20)))
    _insert_ordered(paragraph._p.get_or_add_pPr(), tabs)
    add_run(paragraph, "•", size=S.FS_BULLET_MARK, color=S.GOLD)
    add_run(paragraph, "\t", size=S.FS_BODY, color=S.BODY)
    add_run(paragraph, text, size=S.FS_BODY, color=S.BODY)
    return paragraph


def spec_table(doc, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(
        table,
        {
            "top": (S.TABLE_BORDER, S.BORDER_HAIRLINE_SZ),
            "bottom": (S.TABLE_BORDER, S.BORDER_HAIRLINE_SZ),
            "insideH": (S.TABLE_BORDER, S.BORDER_HAIRLINE_SZ),
        },
    )
    for row, (label, value) in zip(table.rows, rows):
        label_cell, value_cell = row.cells
        label_cell.width = Pt(S.TABLE_COL_LABEL)
        value_cell.width = Pt(S.TABLE_COL_VALUE)
        cell_shading(label_cell, S.TABLE_LABEL_BG)
        cell_margins(label_cell, top=5, bottom=5, left=8, right=6)
        cell_margins(value_cell, top=5, bottom=5, left=5, right=6)

        paragraph = label_cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, label, size=S.FS_TABLE, color=S.NAVY, bold=True)

        paragraph = value_cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, value, size=S.FS_TABLE, color=S.BODY)
    return table


def callout(doc, title: str, paragraphs: list[str], marker: str = "\U0001F4CD"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table_borders(
        table,
        {
            edge: (S.GOLD, S.BORDER_CALLOUT_SZ)
            for edge in ("top", "start", "bottom", "end")
        },
    )
    cell = table.cell(0, 0)
    cell.width = Pt(S.CALLOUT_WIDTH)
    cell_shading(cell, S.CALLOUT_BG)
    cell_margins(cell, top=10, bottom=10, left=12, right=12)

    head = cell.paragraphs[0]
    head.paragraph_format.space_after = Pt(6)
    add_run(head, f"{marker}  ", size=S.FS_SECTION, color=S.BODY)
    add_run(head, title, size=S.FS_SECTION, color=S.NAVY, bold=True)

    for index, text in enumerate(paragraphs):
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(
            0 if index == len(paragraphs) - 1 else 6
        )
        add_run(paragraph, text, size=S.FS_BODY, color=S.BODY)
    return table


def picture(doc, path: Path, *, max_width_cm=S.PHOTO_WIDTH_CM, max_height_cm=S.PHOTO_MAX_HEIGHT_CM):
    with Image.open(path) as img:
        width_px, height_px = img.size
    width_cm = max_width_cm
    height_cm = width_cm * height_px / width_px
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * width_px / height_px
    paragraph = add_paragraph(doc, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    return paragraph


def page_break(doc):
    paragraph = add_paragraph(doc, after=0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph


# --------------------------------------------------------------------------
# страницы отчёта
# --------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(S.PAGE_MARGIN_TOP)
    section.bottom_margin = Cm(S.PAGE_MARGIN_BOTTOM)
    section.left_margin = Cm(S.PAGE_MARGIN_LEFT)
    section.right_margin = Cm(S.PAGE_MARGIN_RIGHT)

    normal = doc.styles["Normal"]
    normal.font.name = S.FONT
    normal.font.size = Pt(S.FS_BODY)
    normal.font.color.rgb = RGBColor.from_string(S.BODY)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15
    return doc


def cover_page(doc, report: dict) -> None:
    for _ in range(6):
        add_paragraph(doc, after=0)

    gold_rule(doc, after=14)

    paragraph = add_paragraph(doc, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = add_run(
        paragraph, report["title"], size=S.FS_COVER_TITLE, color=S.NAVY, bold=True
    )
    run_tracking(run, S.TRACKING_COVER_TITLE)

    paragraph = add_paragraph(doc, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(
        paragraph,
        report["subtitle"],
        size=S.FS_COVER_SUBTITLE,
        color=S.MUTED,
        italic=True,
    )

    gold_rule(doc, after=22, size=S.RULE_NAVY_SZ + 4)

    paragraph = add_paragraph(doc, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(
        paragraph,
        report.get("intro", "В подборку входят:"),
        size=S.FS_COVER_INTRO,
        color=S.MUTED,
    )

    for index, obj in enumerate(report["objects"], start=1):
        paragraph = add_paragraph(doc, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(paragraph, f"{index}. ", size=S.FS_COVER_ITEM, color=S.GOLD, bold=True)
        add_run(
            paragraph,
            f"{obj['street']}, ",
            size=S.FS_COVER_ITEM,
            color=S.NAVY,
            bold=True,
        )
        add_run(
            paragraph,
            f"{obj['city']} — {obj['price_short']}",
            size=S.FS_COVER_ITEM,
            color=S.BODY,
        )

    if report.get("footnote"):
        add_paragraph(doc, after=0)
        paragraph = add_paragraph(doc, before=18, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(paragraph, report["footnote"], size=9, color=S.MUTED, italic=True)


def object_pages(doc, obj: dict, images: list[Path]) -> None:
    gold_rule(doc, after=10)

    paragraph = add_paragraph(doc, after=4)
    add_run(paragraph, obj["title"], size=S.FS_OBJECT_TITLE, color=S.NAVY, bold=True)

    paragraph = add_paragraph(doc, after=12)
    add_run(
        paragraph,
        obj["subtitle"],
        size=S.FS_OBJECT_SUBTITLE,
        color=S.MUTED,
        italic=True,
    )

    spec_table(doc, obj["specs"])
    add_paragraph(doc, after=0)

    for block in obj["sections"]:
        kind = block.get("type", "bullets")
        if block.get("heading"):
            section_heading(doc, block["heading"])
        if kind == "callout":
            callout(doc, block["title"], block["paragraphs"])
        elif kind == "paragraphs":
            for text in block["paragraphs"]:
                body_paragraph(doc, text)
        else:
            for item in block["bullets"]:
                bullet(doc, item)

    if images:
        page_break(doc)
        for index, path in enumerate(images):
            picture(doc, path)
            if index == len(images) - 1:
                break
            if (index + 1) % S.PHOTOS_PER_PAGE == 0:
                page_break(doc)
            else:
                add_paragraph(doc, after=0, before=14)


def build(report: dict, objects: list[tuple[dict, list[Path]]], dest: Path) -> Path:
    layout = report.get("layout")
    if layout == "editorial":
        from . import editorial  # локальный импорт: editorial тянет помощники отсюда

        return editorial.build(report, objects, dest)
    if layout == "dossier":
        from . import dossier

        return dossier.build(report, objects, dest)

    doc = setup_document()
    cover_page(doc, report)
    for obj, images in objects:
        page_break(doc)
        object_pages(doc, obj, images)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    log.info("сохранён DOCX: %s", dest)
    return dest
