"""Рендер отчёта в журнальной схеме (layout: "editorial").

Схема повторяет приёмы, снятые с эталонных изданий (см. style_editorial):
формат 200 × 265 мм вместо A4, тёплая бумага вместо белого листа, основной
кегль 9 pt при интерлиньяже ×1,26, разрыв между крупнейшим и основным
кеглем примерно ×6–8, асимметричная сетка «боковая колонка + основная»,
иллюстрации навылет и колонтитул с нумерацией.

Старая схема (docx_render) остаётся рабочей: обе выбираются полем "layout"
в JSON отчёта, так что оформление можно сравнивать на одних и тех же данных.
"""

from __future__ import annotations

import copy
import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor
from PIL import Image

from . import style_editorial as S
from .docx_render import (
    _el,
    _get_or_add,
    _insert_ordered,
    cell_margins,
    cell_shading,
    keep_with_next,
    paragraph_border,
    run_tracking,
    table_borders,
)

log = logging.getLogger(__name__)


# --------------------------------------------------------------------------
# текст
# --------------------------------------------------------------------------
def txt(
    paragraph,
    text: str,
    *,
    font: str = S.SANS,
    size: float = S.FS_BODY,
    color: str = S.BODY,
    bold: bool = False,
    italic: bool = False,
    tracking: int | None = None,
    caps: bool = False,
):
    run = paragraph.add_run(text.upper() if caps else text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = _get_or_add(run._r.get_or_add_rPr(), "rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), font)
    if tracking:
        run_tracking(run, tracking)
    return run


def par(
    container,
    *,
    before: float = 0,
    after: float = 0,
    align=None,
    lead: float | None = None,
    left: float | None = None,
    right: float | None = None,
    hanging: float | None = None,
):
    paragraph = container.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    if lead is not None:
        fmt.line_spacing = Pt(lead)      # точный интерлиньяж, не множитель
    if left is not None:
        fmt.left_indent = Pt(left)
    if right is not None:
        fmt.right_indent = Pt(right)
    if hanging is not None:
        fmt.first_line_indent = Pt(-hanging)
    return paragraph


def micro(container, text: str, *, color: str = S.BRASS, after: float = 4, before: float = 0):
    """Надстрочная подпись капителью — основной навигационный элемент схемы."""
    paragraph = par(container, before=before, after=after, lead=S.FS_MICRO + 2)
    txt(
        paragraph,
        text,
        font=S.SANS_MEDIUM,
        size=S.FS_MICRO,
        color=color,
        tracking=S.TR_MICRO,
        caps=True,
    )
    return paragraph


def rule(container, *, color: str = S.RULE, size: int = S.SZ_HAIRLINE, before=0, after=0, space=0):
    paragraph = par(container, before=before, after=after, lead=1)
    paragraph_border(paragraph, "bottom", color, size, space=space)
    return paragraph


def spacer(container, mm: float) -> None:
    """Вертикальный отбив точной высоты — набирается кусками, чтобы LibreOffice
    не обрезал слишком длинную строку с фиксированным интерлиньяжем."""
    remaining = mm * 72 / 25.4
    while remaining > 0.5:
        step = min(30.0, remaining)
        par(container, lead=step)
        remaining -= step


def page_break(container) -> None:
    paragraph = par(container, after=0, lead=1)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------------------
# иллюстрации
# --------------------------------------------------------------------------
def _fit(path: Path, width_mm: float, max_height_mm: float) -> tuple[float, float]:
    with Image.open(path) as img:
        w, h = img.size
    height = width_mm * h / w
    if height > max_height_mm:
        height = max_height_mm
        width_mm = height * w / h
    return width_mm, height


def bleed_image(container, path: Path, *, x_mm: float, y_mm: float, w_mm: float, h_mm: float):
    """Изображение навылет: плавающий якорь с привязкой к краю страницы.

    Обычная встроенная картинка не может выйти за поля, поэтому wp:inline
    подменяется на wp:anchor с абсолютными координатами от угла листа.
    """
    paragraph = par(container, after=0, lead=1)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Mm(w_mm), height=Mm(h_mm))
    inline = run._r.find(qn("w:drawing"))[0]

    anchor = OxmlElement("wp:anchor")
    for key, value in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("simplePos", "0"), ("relativeHeight", "1"), ("behindDoc", "1"),
        ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1"),
    ):
        anchor.set(key, value)

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)

    for tag, offset in (("wp:positionH", x_mm), ("wp:positionV", y_mm)):
        node = OxmlElement(tag)
        node.set("relativeFrom", "page")
        value = OxmlElement("wp:posOffset")
        value.text = str(int(Emu(Mm(offset))))
        node.append(value)
        anchor.append(node)

    extent = copy.deepcopy(inline.find(qn("wp:extent")))
    extent.set("cx", str(int(Emu(Mm(w_mm)))))
    extent.set("cy", str(int(Emu(Mm(h_mm)))))
    anchor.append(extent)
    anchor.append(OxmlElement("wp:effectExtent"))
    anchor.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "a:graphic"):
        node = inline.find(qn(tag))
        if node is not None:
            anchor.append(copy.deepcopy(node))
    inline.getparent().replace(inline, anchor)
    return paragraph


def crop_to_ratio(path: Path, ratio: float, cache: Path) -> Path:
    """Центральное кадрирование под единую пропорцию полосы.

    В листингах кадры приходят и 3:2, и 4:3, и вертикальные. Если ставить их
    как есть, ширина иллюстрации гуляет от 88 до 170 мм и полоса рассыпается.
    В эталонных изданиях кадр всегда подрезан под сетку, поэтому и здесь все
    фотографии приводятся к одному отношению сторон.
    """
    cache.mkdir(parents=True, exist_ok=True)
    dest = cache / f"{path.stem}-{ratio:.3f}.jpg"
    if dest.exists():
        return dest
    with Image.open(path) as img:
        img = img.convert("RGB")
        width, height = img.size
        if width / height > ratio:
            new_width = int(round(height * ratio))
            left = (width - new_width) // 2
            box = (left, 0, left + new_width, height)
        else:
            new_height = int(round(width / ratio))
            top = (height - new_height) // 2
            box = (0, top, width, top + new_height)
        img.crop(box).save(dest, quality=88)
    return dest


def photo(container, path: Path, *, width_mm: float = S.CONTENT_W_MM, max_h: float = S.PHOTO_MAX_H_MM):
    width, height = _fit(path, width_mm, max_h)
    # без точного интерлиньяжа: он обрезал бы встроенную картинку до высоты строки
    paragraph = par(container, after=0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.add_run().add_picture(str(path), width=Mm(width), height=Mm(height))
    return paragraph, height


# --------------------------------------------------------------------------
# сетка
# --------------------------------------------------------------------------
def fixed_layout(table, widths_mm: list[float]) -> None:
    """Жёстко фиксирует ширины колонок.

    Одного tcW у ячеек мало: при tblLayout="fixed" ширины берутся из tblGrid,
    который python-docx заполняет равными долями. Без правки сетка «боковая
    колонка + основная» превращалась в две колонки по 85 мм, а длинные слова
    в заголовках рвались посреди слова.
    """
    _insert_ordered(table._tbl.tblPr, _el("tblLayout", type="fixed"))
    _insert_ordered(table._tbl.tblPr, _el("tblW", w=int(sum(widths_mm) * 56.6929), type="dxa"))
    grid_element = table._tbl.find(qn("w:tblGrid"))
    for column, width in zip(grid_element.findall(qn("w:gridCol")), widths_mm):
        column.set(qn("w:w"), str(int(width * 56.6929)))   # мм → двадцатые доли pt


class Flow:
    """Учёт занятой высоты страницы.

    Раздел с заголовком в боковой колонке разрывается некрасиво: подпись
    остаётся внизу страницы, а список уезжает на следующую. Точной вёрстки в
    DOCX нет, поэтому высота блока оценивается по числу строк и блок целиком
    переносится, если в остатке страницы он не помещается.
    """

    HEIGHT_MM = S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM - 18.0
    CHARS_PER_LINE = 72          # 120 мм основной колонки при 9 pt Inter, с запасом
    LINE_MM = S.LH_BODY * 25.4 / 72

    def __init__(self, doc):
        self.doc = doc
        self.used = 0.0

    def new_page(self) -> None:
        page_break(self.doc)
        self.used = 0.0

    def take(self, mm: float) -> None:
        self.used += mm

    def estimate(self, items: list[str], *, extra_mm: float = 0.0) -> float:
        lines = sum(max(1, -(-len(text) // self.CHARS_PER_LINE)) for text in items)
        gaps = len(items) * (5 * 25.4 / 72)
        return lines * self.LINE_MM + gaps + extra_mm

    def fit(self, mm: float, budget: float | None = None) -> bool:
        """Резервирует место; при нехватке начинает новую страницу.

        budget позволяет считать по реальной высоте полосы там, где высота
        блока известна точно (таблицы, иллюстрации), а не по заниженной с
        запасом оценке текста.
        """
        limit = self.HEIGHT_MM if budget is None else budget
        if self.used and self.used + mm > limit:
            page_break(self.doc)
            self.used = mm
            return False
        self.used += mm
        return True

    SHEET_BUDGET_MM = S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM - 6.0


def grid(doc, *, rail_mm: float = S.RAIL_W_MM, gutter_mm: float = S.GUTTER_MM):
    """Асимметричная сетка «боковая колонка — основная» на невидимой таблице.

    Таблица, а не колонки секции: LibreOffice воспроизводит её один в один,
    а колоночные разрывы в DOCX ведут себя непредсказуемо.
    """
    main_mm = S.CONTENT_W_MM - rail_mm - gutter_mm
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {})
    fixed_layout(table, [rail_mm, main_mm + gutter_mm])
    rail, main = table.rows[0].cells
    rail.width = Mm(rail_mm)
    main.width = Mm(main_mm + gutter_mm)
    cell_margins(rail, top=0, bottom=0, left=0, right=gutter_mm * 72 / 25.4)
    cell_margins(main, top=0, bottom=0, left=0, right=0)
    for cell in (rail, main):
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.line_spacing = Pt(1)
    return rail, main


def _clean(cell):
    """Убирает пустой первый абзац ячейки, добавленный python-docx."""
    first = cell.paragraphs[0]
    if not first.runs:
        first._p.getparent().remove(first._p)


# --------------------------------------------------------------------------
# блоки
# --------------------------------------------------------------------------
_NUM = re.compile(
    r"^\s*(?:≈\s*)?(?:€\s*[\d. ]+|[\d.,]+\s*%|[\d., ]+\s*м²|[\d.,]+\s*[а-я]{0,3})"
)


def kpi_split(value: str) -> tuple[str, str]:
    """Делит значение на крупное число и уточнение при нём."""
    match = _NUM.match(value)
    if not match:
        return value, ""
    head = match.group(0).strip()
    tail = value[match.end():].strip(" ,;—-")
    return head, tail


def kpi_band(doc, items: list[tuple[str, str, str]]):
    """Полоса ключевых цифр: подпись капителью, крупное число, уточнение."""
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {"top": (S.BRASS, S.SZ_RULE), "bottom": (S.RULE, S.SZ_HAIRLINE)})
    width = S.CONTENT_W_MM / len(items)
    fixed_layout(table, [width] * len(items))
    for cell, (label, big, small) in zip(table.rows[0].cells, items):
        cell.width = Mm(width)
        cell_margins(cell, top=7, bottom=8, left=0, right=5)
        head = cell.paragraphs[0]
        head.paragraph_format.space_after = Pt(3)
        head.paragraph_format.line_spacing = Pt(S.FS_MICRO + 2)
        txt(head, label, font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.MUTED,
            tracking=S.TR_MICRO, caps=True)

        value = cell.add_paragraph()
        value.paragraph_format.space_after = Pt(2)
        value.paragraph_format.line_spacing = Pt(S.FS_NUMERAL + 1)
        txt(value, big, font=S.SANS_LIGHT, size=S.FS_NUMERAL, color=S.INK, tracking=-6)

        if small:
            note = cell.add_paragraph()
            note.paragraph_format.space_after = Pt(0)
            note.paragraph_format.line_spacing = Pt(S.LH_SMALL)
            txt(note, small, size=S.FS_CAPTION, color=S.MUTED)
    return table


def _group_of(label: str) -> int:
    low = label.lower()
    for index, (_, keys) in enumerate(S.SPEC_GROUPS):
        if any(key in low for key in keys):
            return index
    return len(S.SPEC_GROUPS)


def spec_sheet(doc, specs: list[list[str]], flow: "Flow | None" = None):
    """Характеристики, разложенные по смысловым группам.

    Шестнадцать строк «подпись — значение» подряд читаются как выгрузка из
    базы; в эталонных изданиях данные всегда сгруппированы.
    """
    names = [name for name, _ in S.SPEC_GROUPS] + ["Прочее"]
    buckets: dict[int, list[list[str]]] = {}
    links: list[str] = []
    for row in specs:
        # ссылка на листинг уходит из таблицы в сноску: голый URL в колонке
        # значений и выглядит инородно, и оставляет висячую строку на полосе
        if str(row[1]).startswith("http"):
            links.append(str(row[1]))
            continue
        buckets.setdefault(_group_of(row[0]), []).append(row)

    label_mm = 46.0
    for index in sorted(buckets):
        rows = buckets[index]
        # короткую хвостовую группу не переносим: одна-две строки на отдельной
        # полосе выглядят хуже, чем те же строки внизу заполненной полосы
        if flow is not None and len(rows) > 2:
            # ширина колонки значений — 124 мм при 8 pt, отсюда своя оценка
            values = [f"{label} {value}" for label, value in rows]
            flow.fit(flow.estimate(values, extra_mm=8) * 0.58,
                     budget=flow.SHEET_BUDGET_MM)
        micro(doc, names[index], before=9, after=3)
        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table_borders(
            table,
            {
                "top": (S.RULE, S.SZ_HAIRLINE),
                "bottom": (S.RULE, S.SZ_HAIRLINE),
                "insideH": (S.RULE_SOFT, S.SZ_HAIRLINE),
            },
        )
        fixed_layout(table, [label_mm, S.CONTENT_W_MM - label_mm])
        for row, (label, value) in zip(table.rows, rows):
            label_cell, value_cell = row.cells
            label_cell.width = Mm(label_mm)
            value_cell.width = Mm(S.CONTENT_W_MM - label_mm)
            cell_margins(label_cell, top=3.4, bottom=3.4, left=0, right=8)
            cell_margins(value_cell, top=3.4, bottom=3.4, left=0, right=0)

            head = label_cell.paragraphs[0]
            head.paragraph_format.space_after = Pt(0)
            head.paragraph_format.line_spacing = Pt(S.LH_SMALL)
            txt(head, label, font=S.SANS_MEDIUM, size=S.FS_SMALL, color=S.INK)

            body = value_cell.paragraphs[0]
            body.paragraph_format.space_after = Pt(0)
            body.paragraph_format.line_spacing = Pt(S.LH_SMALL)
            txt(body, value, font=S.SANS, size=S.FS_SMALL, color=S.BODY)
    if index == max(buckets) and links:
        source_note(doc, links)


def source_note(doc, links: list[str]) -> None:
    """Сноска с адресом публикации — после таблицы, кеглем подписи."""
    for link in links:
        paragraph = par(doc, before=6, after=0, lead=S.LH_SMALL)
        txt(paragraph, "Публикация объекта: ", font=S.SANS_MEDIUM,
            size=S.FS_CAPTION, color=S.MUTED)
        txt(paragraph, link, size=S.FS_CAPTION, color=S.MUTED)


def panel(doc, label: str, title: str, paragraphs: list[str]):
    """Врезка «О городе и районе»: плашка на тон глубже бумаги, латунная
    линейка сверху, без рамки и без пиктограммы."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {"top": (S.BRASS, S.SZ_ACCENT)})
    fixed_layout(table, [S.CONTENT_W_MM])
    cell = table.cell(0, 0)
    cell.width = Mm(S.CONTENT_W_MM)
    cell_shading(cell, S.PANEL)
    cell_margins(cell, top=10, bottom=11, left=12, right=12)

    head = cell.paragraphs[0]
    head.paragraph_format.space_after = Pt(2)
    head.paragraph_format.line_spacing = Pt(S.FS_MICRO + 2)
    txt(head, label, font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.BRASS,
        tracking=S.TR_MICRO, caps=True)

    name = cell.add_paragraph()
    name.paragraph_format.space_after = Pt(7)
    name.paragraph_format.line_spacing = Pt(S.FS_H2 + 3)
    txt(name, title, font=S.SERIF_MEDIUM, size=S.FS_H2, color=S.INK)

    for i, text in enumerate(paragraphs):
        body = cell.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body.paragraph_format.space_after = Pt(0 if i == len(paragraphs) - 1 else 6)
        body.paragraph_format.line_spacing = Pt(S.LH_BODY)
        txt(body, text, size=S.FS_BODY, color=S.BODY)
    return table


def listing(cell, items: list[str]):
    """Список без круглых маркеров: латунное тире и висячий отступ."""
    for item in items:
        paragraph = par(cell, after=5, lead=S.LH_BODY, left=13, hanging=13)
        txt(paragraph, "—", font=S.SANS, size=S.FS_BODY, color=S.BRASS)
        txt(paragraph, " ", size=S.FS_BODY, color=S.BODY)
        txt(paragraph, item, size=S.FS_BODY, color=S.BODY)


def railed_block(doc, heading: str, render, *, before: float = 13):
    """Раздел с заголовком в боковой колонке — базовый разворот схемы."""
    par(doc, after=0, lead=before)
    rail, main = grid(doc)
    _clean(rail)
    _clean(main)
    micro(rail, heading, after=0)
    render(main)


# --------------------------------------------------------------------------
# страницы
# --------------------------------------------------------------------------
def _page_setup(section) -> None:
    section.page_width = Mm(S.PAGE_W_MM)
    section.page_height = Mm(S.PAGE_H_MM)
    section.top_margin = Mm(S.MARGIN_TOP_MM)
    section.bottom_margin = Mm(S.MARGIN_BOTTOM_MM)
    section.left_margin = Mm(S.MARGIN_LEFT_MM)
    section.right_margin = Mm(S.MARGIN_RIGHT_MM)
    section.footer_distance = Mm(8.0)


def setup_document() -> Document:
    doc = Document()
    _page_setup(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.size = Pt(S.FS_BODY)
    normal.font.color.rgb = RGBColor.from_string(S.BODY)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(S.LH_BODY)
    fonts = _get_or_add(normal.element.get_or_add_rPr(), "rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), S.SANS)

    # тёплая бумага вместо белого листа — самый заметный признак издания
    background = OxmlElement("w:background")
    background.set(qn("w:color"), S.PAPER)
    doc.element.insert(0, background)
    doc.settings.element.insert(0, OxmlElement("w:displayBackgroundShape"))
    return doc


def _field(paragraph, instruction: str, *, size: float, color: str, font: str = S.SANS):
    """Поле Word (например PAGE) — нумерация должна считаться при печати."""
    for kind, text in (("begin", None), (None, instruction), ("end", None)):
        run = paragraph.add_run()
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        fonts = _get_or_add(run._r.get_or_add_rPr(), "rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            fonts.set(qn(f"w:{attr}"), font)
        if kind:
            node = _el("fldChar", fldCharType=kind)
        else:
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
        run._r.append(node)


def running_footer(section, left_text: str) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(S.FS_MICRO + 2)
    paragraph_border(paragraph, "top", S.RULE, S.SZ_HAIRLINE, space=6)
    tabs = _el("tabs")
    tabs.append(_el("tab", val="right", pos=int(S.CONTENT_W_MM * 72 / 25.4 * 20)))
    _insert_ordered(paragraph._p.get_or_add_pPr(), tabs)
    txt(paragraph, left_text, font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.MUTED,
        tracking=S.TR_EYEBROW, caps=True)
    txt(paragraph, "\t", size=S.FS_MICRO, color=S.MUTED)
    _field(paragraph, " PAGE ", size=S.FS_CAPTION, color=S.INK, font=S.SANS_MEDIUM)


def cover(doc, report: dict, image: Path | None) -> None:
    if image:
        bleed_image(doc, image, x_mm=0, y_mm=0, w_mm=S.PAGE_W_MM, h_mm=S.COVER_PHOTO_H_MM)
        spacer(doc, S.COVER_PHOTO_H_MM - S.MARGIN_TOP_MM + 12)
    else:
        spacer(doc, 60)

    rule(doc, color=S.BRASS, size=S.SZ_ACCENT, after=9)
    eyebrow = par(doc, after=10, lead=S.FS_MICRO + 2)
    txt(eyebrow, report.get("eyebrow", "Инвестиционная подборка · Нидерланды"),
        font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.BRASS,
        tracking=S.TR_EYEBROW, caps=True)

    heading = par(doc, after=8, lead=S.LH_DISPLAY)
    txt(heading, report["title"], font=S.SERIF_LIGHT, size=S.FS_DISPLAY,
        color=S.INK, tracking=S.TR_DISPLAY)

    sub = par(doc, after=0, lead=S.LH_LEAD, right=90)
    txt(sub, report["subtitle"], font=S.SERIF_LIGHT, size=S.FS_LEAD, color=S.BODY)


def cover_footer(doc, report: dict, objects: list[tuple[dict, list[Path]]]) -> None:
    """Нижняя строка обложки — один абзац с правой табуляцией.

    Таблица здесь добавляла высоту и выталкивала обложку на вторую страницу.
    """
    rule(doc, before=13, after=6)
    line = par(doc, after=0, lead=S.LH_SMALL)
    tabs = _el("tabs")
    tabs.append(_el("tab", val="right", pos=int(S.CONTENT_W_MM * 72 / 25.4 * 20)))
    _insert_ordered(line._p.get_or_add_pPr(), tabs)
    txt(line, report.get("source_line", "funda in business"), size=S.FS_CAPTION, color=S.MUTED)
    txt(line, "\t", size=S.FS_CAPTION, color=S.MUTED)
    txt(line, f"{len(objects)} объекта", font=S.SANS_MEDIUM, size=S.FS_CAPTION,
        color=S.INK, tracking=S.TR_MICRO, caps=True)


def overview(doc, report: dict, objects: list[tuple[dict, list[Path]]]) -> None:
    micro(doc, "Обзор подборки", after=7)
    heading = par(doc, after=9, lead=S.LH_H1)
    txt(heading, "Три объекта", font=S.SERIF_LIGHT, size=S.FS_H1, color=S.INK,
        tracking=S.TR_H1)

    if report.get("intro"):
        lead = par(doc, after=15, lead=S.LH_LEAD, right=100)
        txt(lead, report["intro"], font=S.SERIF_LIGHT, size=S.FS_LEAD, color=S.BODY)

    totals = report.get("totals")
    if totals:
        kpi_band(doc, [(label, big, small) for label, big, small in totals])

    micro(doc, "Сравнение", before=15, after=5)
    header = ("", "Объект", "Цена", "Доход / год", "BAR", "Площадь", "Трасса")
    widths = (10.0, 47.0, 24.0, 25.0, 13.0, 21.0, 30.0)
    table = doc.add_table(rows=len(objects) + 1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(
        table,
        {
            "top": (S.INK, S.SZ_RULE),
            "bottom": (S.INK, S.SZ_RULE),
            "insideH": (S.RULE, S.SZ_HAIRLINE),
        },
    )
    fixed_layout(table, list(widths))
    for cell, label, width in zip(table.rows[0].cells, header, widths):
        cell.width = Mm(width)
        cell_margins(cell, top=5, bottom=5, left=0, right=4)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(S.FS_MICRO + 2)
        txt(paragraph, label, font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.MUTED,
            tracking=S.TR_H3, caps=True)

    for index, (obj, _) in enumerate(objects, start=1):
        row = table.rows[index]
        summary = obj.get("summary", {})
        values = (
            f"{index:02d}",
            obj["street"],
            obj.get("price_short", ""),
            summary.get("rent", "—"),
            summary.get("yield", "—"),
            summary.get("area", "—"),
            summary.get("road", "—"),
        )
        for position, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            cell.width = Mm(width)
            cell_margins(cell, top=6, bottom=6, left=0, right=4)
            if index % 2 == 0:
                cell_shading(cell, S.PANEL_SOFT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(S.LH_SMALL)
            if position == 0:
                txt(paragraph, value, font=S.SANS_LIGHT, size=S.FS_SMALL + 2, color=S.BRASS)
            elif position == 1:
                txt(paragraph, value, font=S.SANS_MEDIUM, size=S.FS_SMALL, color=S.INK)
                city = cell.add_paragraph()
                city.paragraph_format.space_after = Pt(0)
                city.paragraph_format.line_spacing = Pt(S.LH_SMALL)
                txt(city, obj.get("city", ""), size=S.FS_CAPTION, color=S.MUTED)
            else:
                txt(paragraph, value, size=S.FS_SMALL, color=S.BODY)

    if report.get("criteria"):
        def render(main, items=report["criteria"]):
            listing(main, items)

        railed_block(doc, "Что искали", render, before=17)

    if report.get("rejected"):
        def render(main, text=report["rejected"]):
            paragraph = par(main, after=0, lead=S.LH_BODY, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            txt(paragraph, text, size=S.FS_BODY, color=S.BODY)

        railed_block(doc, "Что не прошло", render, before=13)


def _opener_image(images: list[Path], index: int) -> Path | None:
    """Кадр шмуцтитула. У первого объекта первый кадр уже занят обложкой."""
    wanted = 2 if index == 1 else 1
    for position in (wanted, 1, 0):
        if len(images) > position:
            return images[position]
    return None


def object_opener(doc, index: int, obj: dict, image: Path | None) -> None:
    if image:
        bleed_image(doc, image, x_mm=0, y_mm=0, w_mm=S.PAGE_W_MM, h_mm=S.OPENER_PHOTO_H_MM)
        spacer(doc, S.OPENER_PHOTO_H_MM - S.MARGIN_TOP_MM + 12)

    rail, main = grid(doc, rail_mm=38.0)
    _clean(rail)
    _clean(main)

    ordinal = par(rail, after=0, lead=S.FS_ORDINAL * 0.78)
    txt(ordinal, f"{index:02d}", font=S.SERIF_LIGHT, size=S.FS_ORDINAL,
        color=S.BRASS, tracking=-30)

    micro(main, obj.get("city", ""), after=6)
    title = par(main, after=7, lead=S.LH_H1)
    txt(title, obj["title"], font=S.SERIF_LIGHT, size=S.FS_H1, color=S.INK,
        tracking=S.TR_H1)
    sub = par(main, after=0, lead=S.LH_SMALL)
    txt(sub, obj["subtitle"], size=S.FS_SMALL, color=S.MUTED)


def object_brief(doc, obj: dict, flow: Flow) -> None:
    micro(doc, "Позиция", after=7)
    lead_text = obj.get("lead") or _first_sentence(obj)
    lead = par(doc, after=13, lead=S.LH_LEAD)
    txt(lead, lead_text, font=S.SERIF_LIGHT, size=S.FS_LEAD, color=S.INK)
    flow.take(flow.estimate([lead_text], extra_mm=11) * S.LH_LEAD / S.LH_BODY)

    kpi_band(doc, obj_kpi(obj))
    flow.take(26)

    for block in obj["sections"]:
        if block.get("type") == "callout":
            texts = block["paragraphs"]
            par(doc, after=0, lead=11)
            # врезка залита плашкой: разрыв между страницами разрезал бы её,
            # поэтому она либо помещается целиком, либо уходит на новую полосу
            flow.fit(flow.estimate(texts, extra_mm=26) * 0.92)
            panel(doc, block.get("heading", "О городе и районе"), block["title"], texts)
            break

    for block in obj["sections"]:
        if block.get("type") == "paragraphs":
            texts = block["paragraphs"]

            def render(main, texts=texts):
                for i, text in enumerate(texts):
                    paragraph = par(main, after=0 if i == len(texts) - 1 else 7,
                                    lead=S.LH_BODY, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                    txt(paragraph, text, size=S.FS_BODY, color=S.BODY)

            railed_block(doc, block.get("heading", "Описание"), render)
            break


def _first_sentence(obj: dict) -> str:
    for block in obj["sections"]:
        if block.get("type") == "paragraphs" and block.get("paragraphs"):
            text = block["paragraphs"][0]
            cut = text.find(". ")
            return text if cut < 0 else text[: cut + 1]
    return obj.get("subtitle", "")


def obj_kpi(obj: dict) -> list[tuple[str, str, str]]:
    if obj.get("kpi"):
        return [tuple(item) for item in obj["kpi"]]
    wanted = ("Цена", "Доход от аренды", "Доходность", "Общая площадь")
    found: list[tuple[str, str, str]] = []
    for name in wanted:
        for label, value in obj["specs"]:
            if label.lower().startswith(name.lower()):
                big, small = kpi_split(value)
                found.append((label, big, small))
                break
    return found[:4]


def object_passport(doc, obj: dict, flow: Flow) -> None:
    flow.new_page()
    micro(doc, "Паспорт объекта", after=6)
    heading = par(doc, after=4, lead=S.FS_H2 + 4)
    txt(heading, obj["street"], font=S.SERIF_MEDIUM, size=S.FS_H2, color=S.INK)
    rule(doc, color=S.INK, size=S.SZ_RULE, after=2)
    spec_sheet(doc, obj["specs"], flow)


def object_analysis(doc, obj: dict, flow: Flow) -> None:
    first = True
    for block in obj["sections"]:
        if block.get("type") not in (None, "bullets"):
            continue
        items = block.get("bullets") or []
        if not items:
            continue
        if first:
            flow.new_page()
            micro(doc, "Разбор", after=6)
            heading = par(doc, after=4, lead=S.FS_H2 + 4)
            txt(heading, "Что стоит за цифрами", font=S.SERIF_MEDIUM,
                size=S.FS_H2, color=S.INK)
            rule(doc, color=S.INK, size=S.SZ_RULE, after=0)
            first = False
        # блок целиком переносится на новую страницу, если не помещается:
        # заголовок в боковой колонке не должен отрываться от своего списка
        flow.fit(flow.estimate(items, extra_mm=11))
        railed_block(doc, block.get("heading", ""),
                     lambda main, items=items: listing(main, items))


def photo_pages(doc, images: list[Path], flow: Flow, cache: Path) -> None:
    """Полосы иллюстраций: карта, затем кадры объекта по два на полосу.

    Карта идёт в своей пропорции, фотографии — подрезанными под единый кадр,
    чтобы ширина иллюстрации не гуляла от полосы к полосе.
    """
    if not images:
        return
    flow.new_page()
    micro(doc, "Объект в кадре", after=6)

    # у иллюстраций свой бюджет: текстовая оценка Flow намеренно занижена
    # запасом, а кадр занимает ровно ту высоту, которую мы ему задали
    budget = S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM - 4.0
    height = S.CONTENT_W_MM / S.PHOTO_RATIO
    used = 8.0

    for index, path in enumerate(images):
        if index == 0:
            _, map_height = photo(doc, path)                   # карта — в своей пропорции
            caption = par(doc, before=4, after=0, lead=S.LH_SMALL)
            txt(caption, "Расположение объекта · картографические данные © Google",
                size=S.FS_CAPTION, color=S.MUTED)
            used += map_height + 8
            continue
        if used + S.PHOTO_GAP_MM + height > budget:
            page_break(doc)
            used = 0.0
        else:
            par(doc, after=0, lead=S.PHOTO_GAP_MM * 72 / 25.4)
            used += S.PHOTO_GAP_MM
        photo(doc, crop_to_ratio(path, S.PHOTO_RATIO, cache), max_h=height)
        used += height
    flow.used = used


def colophon(doc, report: dict) -> None:
    micro(doc, "Источники и оговорки", after=7)
    heading = par(doc, after=10, lead=S.FS_H2 + 4)
    txt(heading, "Как читать эту подборку", font=S.SERIF_MEDIUM, size=S.FS_H2, color=S.INK)
    rule(doc, color=S.INK, size=S.SZ_RULE, after=10)

    for label, text in report.get("colophon", []):
        def render(main, body=text):
            paragraph = par(main, after=0, lead=S.LH_BODY, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            txt(paragraph, body, size=S.FS_BODY, color=S.BODY)
        railed_block(doc, label, render, before=11)


# --------------------------------------------------------------------------
# сборка
# --------------------------------------------------------------------------
def build(report: dict, objects: list[tuple[dict, list[Path]]], dest: Path) -> Path:
    doc = setup_document()
    cache = dest.parent.parent / ".cache" / "crops"

    first_photo = next((imgs[1] for _, imgs in objects if len(imgs) > 1), None)
    cover(doc, report, first_photo)
    cover_footer(doc, report, objects)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    _page_setup(body)
    running_footer(body, report.get("running_title", report["title"]))

    overview(doc, report, objects)
    flow = Flow(doc)

    for index, (obj, images) in enumerate(objects, start=1):
        opener = _opener_image(images, index)
        flow.new_page()
        object_opener(doc, index, obj, opener)
        # дальше — сплошной поток: полосы заполняются подряд, разрыв ставится
        # только там, где неделимый блок не помещается в остаток страницы
        flow.new_page()
        object_brief(doc, obj, flow)
        object_passport(doc, obj, flow)
        object_analysis(doc, obj, flow)
        photo_pages(doc, images, flow, cache)

    flow.new_page()
    colophon(doc, report)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    log.info("сохранён DOCX (editorial): %s", dest)
    return dest
