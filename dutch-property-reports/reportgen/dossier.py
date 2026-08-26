"""Схема «dossier» — визуальная, с коротким составом.

Состав задан заказчиком: простая обложка со списком объектов и ценами;
на каждый объект — шмуцтитул, полоса карточек с ключевыми характеристиками,
описание объекта, описание локации и фотографии. Никакого свода подборки,
никакого разбора и никакого колофона.

Полосы, которые должны производить впечатление, собираются как изображения
(модуль visuals): DOCX не умеет ни положить текст на кадр, ни скруглить углы
иллюстрации, ни дать ей тень.

Оформление снято с еженедельного обзора заказчика: формат A4, две колонки по
80 мм со средником 10 мм, Source Serif 4 и Source Sans 3, чернила #16233A и
охра #9C7C38. Текстовые полосы добираются кадром до нижнего поля — высота
считается по тем же величинам, которыми набраны блоки, поэтому полоса
заканчивается на поле, а не белой третью.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from . import maps, style_editorial as S, visuals
from .docx_render import (
    _el,
    _insert_ordered,
    cell_margins,
    cell_shading,
    paragraph_border,
    table_borders,
)
from .editorial import (
    Flow,
    _clean,
    _page_setup,
    bleed_image,
    fixed_layout,
    grid,
    micro,
    page_break,
    par,
    photo,
    rule,
    running_footer,
    setup_document,
    txt,
)

log = logging.getLogger(__name__)

DISPLAY = "Source Serif 4 Light"
CARD_COLUMNS = 3

# Ширина колонки и средник взяты из еженедельника: две полосы по 80 мм с
# промежутком 10 мм. При кегле 9,5 pt это ~42 знака в строке — то, ради чего
# издание вообще держит две колонки: на всю ширину набора строка вдвое длиннее
# комфортной и текст читается как служебная записка.
COL_W_MM = 80.0
COL_GUTTER_MM = S.CONTENT_W_MM - 2 * COL_W_MM   # 10 мм
FACTS_PHOTO_GAP_MM = 11.0
# Предел высоты кадра во всю ширину набора: ниже пропорция уходит в квадрат
# и от исходного снимка остаётся вырезанная середина.
PHOTO_CAP_MM = 150.0
# Карта строится под остаток полосы: ниже MAP_MIN она перестаёт читаться,
# выше MAP_MAX начинает спорить с текстом. MAP_SLACK — воздух до нижнего поля.
MAP_MIN_MM = 62.0
MAP_MAX_MM = 118.0
MAP_SLACK_MM = 4.0
MAP_PX = 1400


# --------------------------------------------------------------------------
# элементы полосы
# --------------------------------------------------------------------------
def running_head(section, left_text: str, right_text: str) -> None:
    """Верхний колонтитул еженедельника: выпуск слева, рубрика охрой справа."""
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(S.FS_MICRO + 3)
    paragraph_border(paragraph, "bottom", S.RULE, S.SZ_HAIRLINE, space=5)
    tabs = _el("tabs")
    tabs.append(_el("tab", val="right", pos=int(S.CONTENT_W_MM * 72 / 25.4 * 20)))
    _insert_ordered(paragraph._p.get_or_add_pPr(), tabs)
    txt(paragraph, left_text, size=S.FS_MICRO, color=S.MUTED)
    txt(paragraph, "\t", size=S.FS_MICRO, color=S.MUTED)
    txt(paragraph, right_text, font=S.SANS_MEDIUM, size=S.FS_MICRO,
        color=S.BRASS, caps=True)


def blank_running(section) -> None:
    """Полоса навылет живёт без колонтитулов: они легли бы прямо на кадр."""
    for area in (section.header, section.footer):
        area.is_linked_to_previous = False
        paragraph = area.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)


def full_bleed(doc, image: Path) -> None:
    bleed_image(doc, image, x_mm=0, y_mm=0, w_mm=S.PAGE_W_MM, h_mm=S.PAGE_H_MM)


def display_title(doc, text: str, *, size: float = 30.0, after: float = 8.0):
    paragraph = par(doc, after=after, lead=size * 1.06)
    txt(paragraph, text, font=DISPLAY, size=size, color=S.INK, tracking=-8)
    return paragraph


# Выделение в наборе: **фрагмент** набирается полужирной антиквой тем же
# цветом. Цветом или подложкой цифры не выделяются — в сплошном тексте это
# читается как ссылка или маркер, а не как акцент издания.
ACCENT = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)


# Число и единица не должны расходиться по строкам: «7,5» в конце одной, «%»
# в начале следующей — обычный дефект автоматической вёрстки.
UNITS = ("%", "м²", "м", "км", "мм", "кг/м²", "A", "Wp", "т", "мин")
NUMBER_UNIT = re.compile(r"(\d)\s+(" + "|".join(re.escape(u) for u in UNITS) + r")(?![\w²])")
CURRENCY = re.compile(r"€\s+(?=[\d])")


def _nbsp(text: str) -> str:
    """Неразрывные пробелы после знака валюты и перед единицами измерения."""
    text = CURRENCY.sub("€\u00a0", text)
    return NUMBER_UNIT.sub(lambda m: f"{m.group(1)}\u00a0{m.group(2)}", text)


def _words(text: str) -> list[str]:
    """Разбивка на слова только по обычным пробелам: неразрывный остаётся в слове."""
    return [word for word in re.split(r"[ \t\n\r]+", text) if word]


def _mark_words(text: str) -> str:
    """Разметка по словам: `**два слова**` → `**два** **слова**`.

    Дальше текст живёт обычной строкой: его меряют, режут по строкам и
    развешивают по колонкам, а пословная разметка переживает любую резку.
    """
    return ACCENT.sub(lambda m: " ".join(f"**{w}**" for w in _words(m.group(1))),
                      _nbsp(text))


def _plain(word: str) -> tuple[str, bool]:
    """Слово без разметки и признак выделения.

    Знаки препинания часто прилипают снаружи (`**%**.`), поэтому звёздочки
    снимаются отовсюду, а не только с краёв слова.
    """
    return word.replace("**", ""), "**" in word


def emphasis(paragraph, text: str, *, size: float = S.FS_BODY,
             color: str = S.BODY, font: str = S.BODY_FONT,
             accent: str = S.SERIF_MEDIUM) -> None:
    """Набор абзаца с выделенными фрагментами."""
    run, bold = [], False
    for word in _words(_mark_words(text)):
        clean, mark = _plain(word)
        if mark != bold and run:
            txt(paragraph, " ".join(run) + " ",
                font=accent if bold else font, size=size, color=color)
            run = []
        bold = mark
        run.append(clean)
    if run:
        txt(paragraph, " ".join(run), font=accent if bold else font,
            size=size, color=color)


def body_paragraphs(container, texts: list[str], *, after: float = 8.0) -> None:
    """Проза антиквой: гротеск на всю полосу читается серо."""
    for position, text in enumerate(texts):
        # выключка влево, а не по формату: без переносов русский текст
        # при выключке по формату разваливается на разреженные строки
        paragraph = par(container, after=0 if position == len(texts) - 1 else after,
                        lead=S.LH_BODY)
        emphasis(paragraph, text)


def _measure_lines(text: str, width_mm: float) -> list[str]:
    """Разбивка абзаца на строки в колонке заданной ширины.

    Считается по тем же метрикам, которыми набирается полоса: Source Serif 4
    в кегле основного текста, выделенные фрагменты — полужирным начертанием.
    LibreOffice переносит чуть иначе, но развеска по колонкам от расхождения
    в одну строку не разваливается.
    """
    faces = {
        False: visuals.font("SourceSerif4-Regular", S.FS_BODY),
        True: visuals.font("SourceSerif4-SemiBold", S.FS_BODY),
    }
    limit = width_mm / 25.4 * visuals.DPI
    lines: list[str] = []
    current: list[str] = []
    width = 0.0
    for word in _words(_mark_words(text)):
        clean, mark = _plain(word)
        step = faces[mark].getlength(clean if not current else " " + clean)
        if current and width + step > limit:
            lines.append(" ".join(current))
            current, width = [word], faces[mark].getlength(clean)
        else:
            current.append(word)
            width += step
    if current:
        lines.append(" ".join(current))
    return lines


# межабзацный отбив в строках — чтобы развеска считала его наравне с текстом
PARAGRAPH_GAP_LINES = 0.6


def _split_evenly(texts: list[str], width_mm: float) -> tuple[list[str], list[str]]:
    """Развеска текста по двум колонкам поровну.

    Резать только по границам абзацев мало: три абзаца в 8, 13 и 11 строк дают
    в лучшем случае 21 против 11. Поэтому абзац при необходимости продолжается
    во второй колонке — как в самом еженедельнике.
    """
    blocks = [_measure_lines(text, width_mm) for text in texts]
    total = sum(len(block) for block in blocks) + PARAGRAPH_GAP_LINES * (len(blocks) - 1)
    target = total / 2

    left: list[str] = []
    right: list[str] = []
    filled = 0.0
    for index, block in enumerate(blocks):
        if index:
            filled += PARAGRAPH_GAP_LINES
        if filled >= target - 0.5:
            right.append(" ".join(block))
            filled += len(block)
            continue
        room = int(round(target - filled))
        if room >= len(block):
            left.append(" ".join(block))
        else:
            # абзац рвётся по строке: хвост уходит в правую колонку
            if room > 0:
                left.append(" ".join(block[:room]))
            right.append(" ".join(block[room:] if room > 0 else block))
        filled += len(block)
    if not left:
        left, right = right[:1], right[1:]
    return left, right


def text_columns(doc, texts: list[str], *, after: float = 8.0):
    """Две колонки по 80 мм со средником 10 мм — набор еженедельника."""
    if len(texts) < 2:
        body_paragraphs(doc, texts, after=after)
        return None
    left_texts, right_texts = _split_evenly(texts, COL_W_MM)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {})
    fixed_layout(table, [COL_W_MM, COL_GUTTER_MM, COL_W_MM])
    left, _, right = table.rows[0].cells
    for cell in (left, right):
        cell.width = Mm(COL_W_MM)
        cell_margins(cell, top=0, bottom=0, left=0, right=0)
    for cell, block in ((left, left_texts), (right, right_texts)):
        first = cell.paragraphs[0]
        first.paragraph_format.space_after = Pt(after if len(block) > 1 else 0)
        first.paragraph_format.line_spacing = Pt(S.LH_BODY)
        emphasis(first, block[0])
        body_paragraphs(cell, block[1:], after=after)
    return table


def lead_paragraph(doc, text: str, *, after: float = 12.0):
    """Лид — курсив антиквы: в еженедельнике это Source Serif 4 Italic 12,5 pt."""
    paragraph = par(doc, after=after, lead=S.LH_LEAD)
    txt(paragraph, text, font=S.SERIF, size=S.FS_LEAD, color=S.INK_SOFT, italic=True)
    return paragraph


def pull_quote(doc, text: str, *, before: float = 13.0, after: float = 13.0):
    """Врезка-цитата курсивом дисплейной антиквы — главный тезис полосы."""
    rule(doc, color=S.BRASS, size=S.SZ_ACCENT, before=before, after=8)
    paragraph = par(doc, after=8, lead=27, right=24)
    txt(paragraph, text, font=S.SERIF, size=22, color=S.INK, italic=True)
    rule(doc, color=S.RULE, size=S.SZ_HAIRLINE, after=after)


PAD_MM = 4.0            # поля ячейки полосы цифр, 6 + 4 pt


# Отбивки перечня характеристик, снятые с готового PDF: считать их из
# space_before/space_after нельзя — соседние отбивки схлопываются, и расчёт
# расходится с вёрсткой на треть полосы.
GROUP_TOP_MM = 2.3       # от отбива до первого заголовка группы
GROUP_GAP_MM = 8.7       # между группами в колонке
GROUP_HEAD_MM = 5.5      # заголовок группы с отбивкой до первой строки
LINK_BLOCK_MM = 14.6     # линейка и строка со ссылкой на публикацию


def _mm(points: float) -> float:
    return points * 25.4 / 72


def _lines(text: str, width_mm: float, face: str, size: float) -> int:
    """Число строк произвольного набора — для расчёта высоты полосы.

    Кегль округляется вниз до половины пункта: DOCX хранит его в половинах
    пункта, поэтому 7,4 pt на полосе выходит семёркой, а не семёркой с
    хвостом, и строка помещается там, где расчёт по 7,4 давал перенос.
    """
    metrics = visuals.font(face, int(size * 2) / 2)
    limit = width_mm / 25.4 * visuals.DPI
    count, current = 1, ""
    for word in str(text).split():
        probe = f"{current} {word}".strip()
        if current and metrics.getlength(probe) > limit:
            count += 1
            current = word
        else:
            current = probe
    return count


def _fit_numerals(values: list[str], width_mm: float) -> float:
    """Наибольший кегль, при котором все значения полосы влезают в строку."""
    for size in (S.FS_NUMERAL, 21, 20, 19, 18, 17, 16, 15, 14):
        face = visuals.font("SourceSerif4-Light", size)
        if all(face.getlength(str(value)) / visuals.DPI * 25.4 <= width_mm
               for value in values):
            return float(size)
    return 14.0


def figures_band(doc, items: list[list[str]]):
    """Полоса главных цифр: цена, доход, доходность, площадь.

    Одна плашка на четыре показателя, а не четыре отдельные карточки: полоса
    задаёт иерархию, с которой начинается чтение, и не спорит с перечнем ниже.
    """
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {
        "top": (S.BRASS, S.SZ_ACCENT),
        "bottom": (S.RULE, S.SZ_HAIRLINE),
        "insideV": (S.RULE, S.SZ_HAIRLINE),
    })
    width = S.CONTENT_W_MM / len(items)
    fixed_layout(table, [width] * len(items))
    # один кегль на всю полосу — крупнейший, при котором самое длинное значение
    # («€ 1.900.000») ещё держится в одну строку. Разнокегельность внутри полосы
    # цифр читается как сбой вёрстки, поэтому размер общий, а не поячеечный
    size = _fit_numerals([entry[1] for entry in items], width - PAD_MM)
    for cell, entry in zip(table.rows[0].cells, items):
        label, value, note = (list(entry) + ["", "", ""])[:3]
        cell.width = Mm(width)
        cell_shading(cell, S.PANEL)
        cell_margins(cell, top=15, bottom=16, left=6, right=4)

        head = cell.paragraphs[0]
        head.paragraph_format.space_after = Pt(4)
        head.paragraph_format.line_spacing = Pt(S.FS_MICRO + 2)
        txt(head, label, font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.BRASS, caps=True)

        big = cell.add_paragraph()
        big.paragraph_format.space_after = Pt(3)
        big.paragraph_format.line_spacing = Pt(size * 1.18)
        txt(big, value, font=S.SERIF_LIGHT, size=size, color=S.BRASS)

        if note:
            caption = cell.add_paragraph()
            caption.paragraph_format.space_after = Pt(0)
            caption.paragraph_format.line_spacing = Pt(9.6)
            txt(caption, note, size=7.4, color=S.MUTED)
    return table


LABEL_INDENT_PT = 96.0          # 34 мм под подпись строки


def fact_row(cell, label: str, value: str) -> None:
    """Строка перечня: подпись слева, значение с висячим отступом справа."""
    paragraph = par(cell, before=5.6, after=5.6, lead=12.8,
                    left=LABEL_INDENT_PT, hanging=LABEL_INDENT_PT)
    tabs = _el("tabs")
    tabs.append(_el("tab", val="left", pos=int(LABEL_INDENT_PT * 20)))
    _insert_ordered(paragraph._p.get_or_add_pPr(), tabs)
    paragraph_border(paragraph, "bottom", S.RULE_SOFT, S.SZ_HAIRLINE, space=3)
    txt(paragraph, label, font=S.SANS_MEDIUM, size=8.8, color=S.INK)
    txt(paragraph, "\t", size=8.8, color=S.BODY)
    txt(paragraph, value, size=8.8, color=S.BODY)


def fact_columns(doc, groups: list) -> None:
    """Характеристики группами в две колонки — перечень, а не сетка плашек.

    Группы раскладываются попарно по строкам таблицы: левая колонка получает
    первую половину, правая — вторую, и заголовок каждой пары стоит на одной
    высоте. Если просто складывать группы в две ячейки подряд, вторые
    заголовки расходятся по вертикали, как только строки в группах
    переносятся по-разному.
    """
    half = -(-len(groups) // 2)
    pairs = [(groups[index], groups[index + half] if index + half < len(groups) else None)
             for index in range(half)]

    column_mm = (S.CONTENT_W_MM - 6.0) / 2
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {})
    fixed_layout(table, [column_mm + 6.0, column_mm])

    for order, (row, pair) in enumerate(zip(table.rows, pairs)):
        for position, cell in enumerate(row.cells):
            cell.width = Mm(column_mm + (6.0 if position == 0 else 0.0))
            cell_margins(cell, top=0, bottom=0, left=0,
                         right=17 if position == 0 else 0)
            _clean(cell)
            group = pair[position]
            if group is None:
                continue
            title, rows = group
            header = par(cell, before=0 if order == 0 else 12, after=5,
                         lead=S.FS_MICRO + 2)
            paragraph_border(header, "top", S.BRASS, S.SZ_RULE, space=5)
            txt(header, title, font=S.SANS_MEDIUM, size=S.FS_MICRO,
                color=S.BRASS, caps=True)
            for label, value in rows:
                fact_row(cell, label, value)


def access_block(doc, title: str, rows: list[list[str]]) -> None:
    """Перечень доступности под картой, разложенный на две колонки.

    Заголовок идёт во всю ширину, а строки делятся пополам: у группы из
    восьми пунктов колонка вышла бы вдвое длиннее соседней.
    """
    header = par(doc, before=13, after=5, lead=S.FS_MICRO + 2)
    paragraph_border(header, "top", S.BRASS, S.SZ_RULE, space=5)
    txt(header, title, font=S.SANS_MEDIUM, size=S.FS_MICRO, color=S.BRASS, caps=True)

    half = -(-len(rows) // 2)
    column_mm = (S.CONTENT_W_MM - 6.0) / 2
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {})
    fixed_layout(table, [column_mm + 6.0, column_mm])
    for position, cell in enumerate(table.rows[0].cells):
        cell.width = Mm(column_mm + (6.0 if position == 0 else 0.0))
        cell_margins(cell, top=0, bottom=0, left=0,
                     right=17 if position == 0 else 0)
        _clean(cell)
        for label, value in (rows[:half] if position == 0 else rows[half:]):
            fact_row(cell, label, value)


FRAME_PAD_MM = 8.0      # поле под тень вокруг кадра: 4 мм с каждой стороны


def frame_ratio(width_mm: float, height_mm: float) -> float:
    """Пропорция кадра, при которой готовый файл встанет ровно в высоту.

    Файл с тенью шире и выше самого снимка на поле под тень, поэтому под
    заданную высоту полосы пропорцию надо считать по внутреннему кадру.
    """
    return (width_mm - FRAME_PAD_MM) / max(height_mm - FRAME_PAD_MM, 1.0)


def framed_photo(doc, source: Path, cache: Path, *, width_mm: float,
                 ratio: float = S.PHOTO_RATIO, align=None):
    """Кадр со скруглением и тенью — на полосу ложится уже готовым файлом."""
    # в имя входят пропорция и время правки исходника: один и тот же кадр
    # берётся и панорамой, и полосой, а перерисованная карта не должна
    # подхватываться из кэша по старому имени
    stamp = int(source.stat().st_mtime)
    prepared = visuals.rounded_photo(
        cache / f"{source.stem}-{int(width_mm)}-{ratio:.3f}-{stamp}.jpg", source,
        width_mm=width_mm, ratio=ratio)
    return photo(doc, prepared, width_mm=width_mm, max_h=260.0, align=align)


# --------------------------------------------------------------------------
# полосы объекта
# --------------------------------------------------------------------------
def _facts_height(index: int, obj: dict, title_size: float) -> float:
    """Оценка высоты полосы характеристик в мм.

    Нужна, чтобы закрыть полосу кадром ровно по нижнему полю: точной вёрстки
    в DOCX нет, поэтому высота собирается из тех же величин, которыми набраны
    блоки. Расхождение с готовым PDF — доли миллиметра на блок.
    """
    used = _mm(S.FS_MICRO + 2) + _mm(6)
    used += _mm(title_size * 1.06) * _lines(
        obj["title"], S.CONTENT_W_MM, "SourceSerif4-Light", title_size) + _mm(8)
    if obj.get("lead"):
        used += _mm(S.LH_LEAD) * _lines(
            obj["lead"], S.CONTENT_W_MM, "SourceSerif4-Italic", S.FS_LEAD) + _mm(11)

    facts = obj.get("facts") or {}
    headline = facts.get("headline") or []
    if headline:
        width = S.CONTENT_W_MM / len(headline)
        size = _fit_numerals([entry[1] for entry in headline], width - PAD_MM)
        notes = max(_lines(entry[2], width - PAD_MM, "SourceSans3-Regular", 7.4)
                    if len(entry) > 2 and entry[2] else 0 for entry in headline)
        used += (_mm(15 + S.FS_MICRO + 2 + 4 + size * 1.18 + 3 + 16)
                 + _mm(9.6) * notes)

    if facts.get("groups"):
        used += _mm(15)
        groups = facts["groups"]
        half = -(-len(groups) // 2)
        column_mm = (S.CONTENT_W_MM - 6.0) / 2
        for order in range(half):
            pair = [groups[order]]
            if order + half < len(groups):
                pair.append(groups[order + half])
            # у левой колонки лишние 6 мм ширины съедает её правое поле,
            # так что место под значение в обеих колонках одинаковое
            value_mm = column_mm - _mm(LABEL_INDENT_PT)
            heights = []
            for group in pair:
                # соседние отбивки абзацев LibreOffice схлопывает в одну,
                # поэтому шаг строки перечня — интерлиньяж плюс один отбив
                height = (GROUP_GAP_MM if order else GROUP_TOP_MM) + GROUP_HEAD_MM
                for position_row, (label, value) in enumerate(group[1]):
                    if position_row:
                        height += _mm(5.6)
                    height += _mm(12.8) * _lines(
                        value, value_mm, "SourceSans3-Regular", 8.8)
                heights.append(height)
            used += max(heights)

    if any(str(value).startswith("http") for _, value in obj.get("specs", [])):
        used += LINK_BLOCK_MM
    return used


def object_facts(doc, index: int, obj: dict, cache: Path,
                 closer: Path | None) -> None:
    # длинное название сажаем на кегль поменьше, иначе оно уходит на две строки
    title_size = 31 if len(obj["title"]) <= 26 else 25
    micro(doc, f"Объект {index:02d} · {obj.get('city', '')}", after=6)
    display_title(doc, obj["title"], size=title_size, after=8)
    if obj.get("lead"):
        lead_paragraph(doc, obj["lead"], after=11)

    facts = obj.get("facts") or {}
    if facts.get("headline"):
        figures_band(doc, facts["headline"])
    if facts.get("groups"):
        par(doc, after=0, lead=15)
        fact_columns(doc, facts["groups"])

    # ссылка на публикацию закрывает полосу и заменяет строку в перечне
    link = next((value for label, value in obj.get("specs", [])
                 if str(value).startswith("http")), None)
    if link:
        par(doc, after=0, lead=14)
        rule(doc, color=S.RULE, size=S.SZ_HAIRLINE, after=5)
        source = par(doc, after=0, lead=S.LH_SMALL)
        txt(source, "Публикация объекта: ", font=S.SANS_MEDIUM,
            size=S.FS_CAPTION, color=S.MUTED)
        txt(source, link, size=S.FS_CAPTION, color=S.MUTED)

    # свободную треть полосы закрывает кадр объекта во всю ширину набора:
    # высота считается по остатку, поэтому полоса заканчивается ровно на поле
    free = (S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM
            - _facts_height(index, obj, title_size) - FACTS_PHOTO_GAP_MM)
    if closer and free >= 42.0:
        par(doc, after=0, lead=FACTS_PHOTO_GAP_MM * 72 / 25.4)
        framed_photo(doc, closer, cache, width_mm=S.CONTENT_W_MM,
                     ratio=frame_ratio(S.CONTENT_W_MM, min(free, PHOTO_CAP_MM)))


def _heading_height(title: str, size: float) -> float:
    """Рубрика, заголовок и линейка под ним — шапка любой текстовой полосы."""
    return (_mm(S.FS_MICRO + 2) + _mm(6)
            + _mm(size * 1.06) * _lines(title, S.CONTENT_W_MM,
                                        "SourceSerif4-Light", size) + _mm(8)
            + _mm(12) + 0.5)


def _columns_height(texts: list[str]) -> float:
    """Высота двухколонного набора — по более длинной колонке.

    Ширина колонки берётся с запасом в три процента: LibreOffice переносит
    чуть раньше расчёта, и без запаса полоса иногда не закрывалась кадром, а
    выталкивала его на отдельную страницу.
    """
    if len(texts) < 2:
        blocks = [texts]
    else:
        left, right = _split_evenly(texts, COL_W_MM)
        blocks = [left, right]
    height = 0.0
    for block in blocks:
        column = sum(_mm(S.LH_BODY) * len(_measure_lines(text, COL_W_MM * 0.97))
                     for text in block)
        column += _mm(8) * (len(block) - 1)
        height = max(height, column)
    return height


def object_description(doc, obj: dict, cache: Path, portrait: Path | None) -> None:
    micro(doc, "Описание объекта", after=6)
    display_title(doc, obj["street"], size=26, after=8)
    rule(doc, color=S.INK, size=S.SZ_RULE, after=11)

    texts: list[str] = []
    for block in obj["sections"]:
        if block.get("type") != "paragraphs":
            continue
        texts = block["paragraphs"]
        text_columns(doc, texts)
        break

    used = _heading_height(obj["street"], 26) + _columns_height(texts)
    if obj.get("pull"):
        pull_quote(doc, obj["pull"])
        used += (_mm(13 + 8) + 0.5
                 + _mm(27) * _lines(obj["pull"], (S.CONTENT_W_MM - _mm(24)) * 0.97,
                                    "SourceSerif4-Italic", 22)
                 + _mm(8 + 13) + 0.5)

    free = S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM - used - 6.0
    if portrait and free >= 42.0:
        # кадр добирает полосу до нижнего поля — вместо белой трети внизу
        par(doc, after=0, lead=3)
        framed_photo(doc, portrait, cache, width_mm=S.CONTENT_W_MM,
                     ratio=frame_ratio(S.CONTENT_W_MM, min(free, PHOTO_CAP_MM)))


def _rows_height(rows: list, columns: int = 2) -> float:
    """Высота перечня «подпись — значение», разложенного по колонкам."""
    half = -(-len(rows) // columns) if columns > 1 else len(rows)
    column_mm = (S.CONTENT_W_MM - 6.0) / 2
    value_mm = column_mm - _mm(LABEL_INDENT_PT)
    height = 0.0
    for start in range(0, len(rows), half):
        chunk = rows[start:start + half]
        column = _mm(5.6) * (len(chunk) - 1)
        for label, value in chunk:
            column += _mm(12.8) * _lines(value, value_mm * 0.97,
                                         "SourceSans3-Regular", 8.8)
        height = max(height, column)
    return height


def object_location(doc, obj: dict, cache: Path, assets: Path, *,
                    skip_map: bool = False) -> None:
    block = next((b for b in obj["sections"] if b.get("type") == "callout"), None)
    if block is None:
        return
    micro(doc, "Локация", after=6)
    display_title(doc, block["title"], size=26, after=8)
    rule(doc, color=S.INK, size=S.SZ_RULE, after=11)
    text_columns(doc, block["paragraphs"])

    # карта занимает весь остаток полосы: перечень доступности прижат к
    # нижнему полю, всё, что выше него, отдаётся карте
    used = (_heading_height(block["title"], 26)
            + _columns_height(block["paragraphs"]) + _mm(13)
            + _mm(3 + S.LH_SMALL))
    if obj.get("access"):
        used += _mm(13 + 5 + S.FS_MICRO + 2 + 5) + _rows_height(obj["access"])
    free = S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM - used - MAP_SLACK_MM
    free = min(max(free, MAP_MIN_MM), MAP_MAX_MM)

    # карта рисуется сразу под остаток полосы, поэтому её не приходится
    # подрезать: подрезка срезала бы то, ради чего карта и стоит в отчёте, —
    # метку объекта или центр города
    ratio = (S.CONTENT_W_MM - FRAME_PAD_MM) / (free - FRAME_PAD_MM)
    overview = None if skip_map else maps.for_object(
        obj, assets, width=MAP_PX, height=int(round(MAP_PX / ratio)))
    if overview:
        par(doc, after=0, lead=13)
        framed_photo(doc, overview, cache, width_mm=S.CONTENT_W_MM, ratio=ratio)
        # подпись выключается по центру кадра, а не по левому краю набора
        caption = par(doc, before=3, after=0, lead=S.LH_SMALL,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        txt(caption, "Расположение объекта · картографические данные © Google",
            size=S.FS_CAPTION, color=S.MUTED)

    if obj.get("access"):
        access_block(doc, "Доступность и окружение", obj["access"])


CHAPTER_SLACK_MM = 6.0
CHAPTER_MIN_MM = 40.0


def _pull_height(text: str) -> float:
    return (_mm(13 + 8) + 0.5
            + _mm(27) * _lines(text, (S.CONTENT_W_MM - _mm(24)) * 0.97,
                               "SourceSerif4-Italic", 22)
            + _mm(8 + 13) + 0.5)


def object_chapter(doc, obj: dict, chapter: dict, cache: Path, assets: Path,
                   images: list[Path]) -> None:
    """Дополнительная полоса объекта: текст в две колонки, кадр и перечень.

    Пяти полос схемы хватает, чтобы показать объект. Когда по объекту нужен
    разбор — состав площадей, арендаторы, экономика, окружение, — полосы
    добавляются карточкой в поле ``chapters``. Набираются они тем же
    инструментом, что и остальные, поэтому в отчёте не видно, где кончается
    схема и начинается дополнение.

    Кадры берутся по номерам из общего списка изображений объекта: один
    закрывает полосу во всю ширину, два-три встают рядом.
    """
    micro(doc, chapter.get("kicker", "Разбор"), after=6)
    display_title(doc, chapter["title"], size=26, after=8)
    rule(doc, color=S.INK, size=S.SZ_RULE, after=11)

    texts = chapter.get("paragraphs", [])
    if texts:
        text_columns(doc, texts)
    used = _heading_height(chapter["title"], 26) + _columns_height(texts)

    if chapter.get("pull"):
        pull_quote(doc, chapter["pull"])
        used += _pull_height(chapter["pull"])

    rows = chapter.get("rows") or []
    if rows:
        used += _mm(13 + 5 + S.FS_MICRO + 2 + 5) + _rows_height(rows)

    frames = [images[index] for index in chapter.get("images", [])
              if 0 <= index < len(images)]
    if frames or chapter.get("map"):
        used += _mm(13) + (_mm(3 + S.LH_SMALL) if chapter.get("caption") else 0.0)
        free = (S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM
                - used - CHAPTER_SLACK_MM)
        # поэтажный план нельзя тянуть на всю оставшуюся полосу: кадр режется
        # под заданную высоту, и растянутый план теряет поля с размерами
        height = min(free, chapter.get("image_height", PHOTO_CAP_MM), PHOTO_CAP_MM)

        if chapter.get("map"):
            # карта строится сразу под остаток полосы и не подрезается
            ratio = (S.CONTENT_W_MM - FRAME_PAD_MM) / max(height - FRAME_PAD_MM, 1.0)
            conf = dict(obj.get("map") or {}, **chapter["map"])
            overview = maps.for_object({**obj, "map": conf}, assets, width=MAP_PX,
                                       height=int(round(MAP_PX / ratio)))
            frames = ([overview] if overview else []) + frames

        if frames and free >= CHAPTER_MIN_MM:
            par(doc, after=0, lead=13)
            photo_row(doc, frames[:4], cache, height)
            if chapter.get("caption"):
                caption = par(doc, before=3, after=0, lead=S.LH_SMALL,
                              align=WD_ALIGN_PARAGRAPH.CENTER)
                txt(caption, chapter["caption"], size=S.FS_CAPTION, color=S.MUTED)

    if rows:
        access_block(doc, chapter.get("rows_title", "Подробности"), rows)


class Plan:
    """Разбор списка изображений объекта по назначению.

    Список приходит плоским: обзорная карта и кадры листинга. Карта узнаётся
    по имени файла, а не по позиции — так порядок кадров не зависит от того,
    строилась карта в этой сборке или нет.
    """

    def __init__(self, images: list[Path]):
        self.overview: Path | None = None
        self.photos: list[Path] = []
        for path in images:
            if path.suffix == ".png" and "-google" in path.name:
                self.overview = path
            else:
                self.photos.append(path)

    @property
    def hero(self) -> Path | None:
        """Кадр шмуцтитула и миниатюры на обложке — общий вид здания.

        Брокеры ставят его первым в галерее, поэтому берётся первый кадр:
        интерьер склада на шмуцтитуле читается как случайная фотография.
        """
        return self.photos[0] if self.photos else None

    @property
    def portrait(self) -> Path | None:
        """Кадр полосы описания — следующий за общим видом."""
        if len(self.photos) > 1:
            return self.photos[1]
        return self.photos[0] if self.photos else None

    @property
    def closer(self) -> Path | None:
        """Кадр, закрывающий полосу характеристик."""
        return self.photos[2] if len(self.photos) > 2 else None

    @property
    def gallery(self) -> list[Path]:
        """Кадры галереи.

        Кадр шмуцтитула из галереи не исключается: на шмуцтитуле он затемнён,
        обрезан под полосу и закрыт заголовком, а в галерее показан целиком.
        При пяти-семи кадрах на объект отдавать один целиком под шмуцтитул
        слишком расточительно.
        """
        used = {self.portrait, self.closer}
        return [path for path in self.photos if path not in used]


GALLERY_GUTTER_MM = 6.0
GALLERY_GAP_MM = 9.0
GALLERY_SLACK_MM = 4.0      # запас до нижнего поля, иначе разрыв даёт пустую полосу


def row_width(count: int) -> float:
    """Ширина одного кадра в ряду из ``count`` штук."""
    return (S.CONTENT_W_MM - GALLERY_GUTTER_MM * (count - 1)) / count


# «Естественная» высота ряда: кадр во всю ширину идёт под 16:10, кадры в ряду
# по два и по три — под 4:3. От этих величин считается растяжка полосы.
NATURAL_MM = {
    1: (S.CONTENT_W_MM - FRAME_PAD_MM) / (16 / 10) + FRAME_PAD_MM,
    2: (row_width(2) - FRAME_PAD_MM) / (4 / 3) + FRAME_PAD_MM,
    3: (row_width(3) - FRAME_PAD_MM) / (4 / 3) + FRAME_PAD_MM,
}

# Раскладки полосы: крупный кадр в связке с рядами мелких. Наборы подобраны
# так, чтобы сумма «естественных» высот была близка к высоте полосы — тогда
# растяжка до нижнего поля не уводит пропорцию кадра в квадрат. На каждое
# число кадров есть несколько вариантов: они чередуются, чтобы соседние
# полосы галереи не повторяли друг друга ритмом.
GALLERY_ROWS: dict[int, tuple[tuple[int, ...], ...]] = {
    6: ((1, 2, 3), (3, 1, 2), (2, 3, 1), (1, 3, 2)),
    5: ((1, 2, 2), (2, 1, 2), (2, 2, 1)),
    4: ((1, 2, 1), (2, 1, 1), (1, 1, 2)),
    3: ((1, 2), (2, 1)),
    2: ((1, 1),),
    1: ((1,),),
}


def _gallery_plan(count: int) -> list[int]:
    """Разбивка кадров по полосам: до шести, без одинокого кадра в хвосте."""
    plan = []
    while count > 0:
        if count in (7, 8):
            take = count - 4      # 7 → 3 + 4, 8 → 4 + 4: ровнее, чем 6 + 1|2
        else:
            take = min(6, count)
        plan.append(take)
        count -= take
    return plan


def photo_row(doc, paths: list[Path], cache: Path, height_mm: float) -> None:
    """Ряд галереи: один кадр во всю ширину набора либо два-три в строку."""
    if len(paths) == 1:
        framed_photo(doc, paths[0], cache, width_mm=S.CONTENT_W_MM,
                     ratio=frame_ratio(S.CONTENT_W_MM, height_mm))
        return
    width = row_width(len(paths))
    widths: list[float] = []
    for position in range(len(paths)):
        if position:
            widths.append(GALLERY_GUTTER_MM)
        widths.append(width)
    table = doc.add_table(rows=1, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_borders(table, {})
    fixed_layout(table, widths)
    cells = table.rows[0].cells
    for index, path in enumerate(paths):
        cell = cells[index * 2]
        cell.width = Mm(width)
        cell_margins(cell, top=0, bottom=0, left=0, right=0)
        _clean(cell)
        framed_photo(cell, path, cache, width_mm=width,
                     ratio=frame_ratio(width, height_mm))


def photo_pages(doc, gallery: list[Path], flow: Flow, cache: Path) -> None:
    """Полосы иллюстраций мозаикой: крупный кадр и ряды мелких.

    Высоты рядов растягиваются до нижнего поля, поэтому полоса заполнена
    целиком, а не обрывается на середине белым полем.
    """
    if not gallery:
        return
    flow.new_page()
    micro(doc, "Объект в кадре", after=7)

    head_mm = _mm(S.FS_MICRO + 2 + 7)
    position = 0
    for order, take in enumerate(_gallery_plan(len(gallery))):
        if order:
            page_break(doc)
        variants = GALLERY_ROWS[take]
        rows = variants[order % len(variants)]
        budget = (S.PAGE_H_MM - S.MARGIN_TOP_MM - S.MARGIN_BOTTOM_MM
                  - (head_mm if order == 0 else 0.0) - GALLERY_SLACK_MM
                  - GALLERY_GAP_MM * (len(rows) - 1))
        natural = [NATURAL_MM[size] for size in rows]
        scale = min(budget / sum(natural), 1.45)
        for step, size in enumerate(rows):
            if step:
                par(doc, after=0, lead=GALLERY_GAP_MM * 72 / 25.4)
            photo_row(doc, gallery[position:position + size], cache,
                      natural[step] * scale)
            position += size
    flow.used = S.PAGE_H_MM


# --------------------------------------------------------------------------
# сборка
# --------------------------------------------------------------------------
def build(report: dict, objects: list[tuple[dict, list[Path]]], dest: Path) -> Path:
    doc = setup_document()
    root = dest.parent.parent
    assets = root / ".cache" / "visuals"
    cache = root / ".cache" / "framed"

    items = []
    for obj, images in objects:
        thumb = Plan(images).hero or images[0]
        items.append((obj["street"], obj.get("city", ""),
                      obj.get("price_short", ""), thumb))
    full_bleed(doc, visuals.contents_cover(
        assets / "cover.jpg",
        kicker=report.get("eyebrow", "Инвестиционная подборка · Нидерланды"),
        title=report.get("cover_title", report["title"]),
        subtitle=report["subtitle"],
        items=items,
        meta=report.get("source_line", ""),
    ))

    def plate_section():
        """Полоса-шмуцтитул: своя секция, чтобы снять с неё колонтитулы."""
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _page_setup(section)
        blank_running(section)
        return section

    def body_section():
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _page_setup(section)
        running_head(section, report.get("running_title", report["title"]),
                     report.get("rubric", "Подборка объектов"))
        running_footer(section, "")
        return section

    flow = Flow(doc)
    started = False
    for index, (obj, images) in enumerate(objects, start=1):
        plan = Plan(images)
        opener_photo = plan.hero
        if opener_photo:
            plate_section()
            full_bleed(doc, visuals.opener(
                assets / f"opener-{index}.jpg", opener_photo,
                ordinal=f"{index:02d}", city=obj.get("city", ""),
                title=obj["title"], subtitle=obj["subtitle"],
                kpi=[(label, value) for label, value, _ in obj.get("kpi", [])[:5]],
            ))
            body_section()
        elif not started:
            body_section()
        else:
            flow.new_page()
        started = True
        object_facts(doc, index, obj, cache, plan.closer)
        flow.new_page()
        object_description(doc, obj, cache, plan.portrait)
        flow.new_page()
        object_location(doc, obj, cache, assets,
                        skip_map=bool(report.get("skip_map")))
        chapters = obj.get("chapters", [])
        for chapter in chapters:
            flow.new_page()
            object_chapter(doc, obj, chapter, cache, assets, plan.photos)
        # кадр, уже показанный в разборе, в галерее не повторяется
        shown = {plan.photos[index] for chapter in chapters
                 for index in chapter.get("images", [])
                 if 0 <= index < len(plan.photos)}
        photo_pages(doc, [p for p in plan.gallery if p not in shown], flow, cache)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    log.info("сохранён DOCX (dossier): %s", dest)
    return dest
