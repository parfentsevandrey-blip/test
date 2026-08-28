#!/usr/bin/env python3
"""Собирает названия фильмов из сохранённых страниц IAFD (person.rme) в один .docx.

Сайт закрыт Cloudflare-челленджем и в robots.txt запрещает краулеры,
поэтому HTML нужно сохранить самому из браузера (Ctrl+S -> "Web Page, HTML only"),
а скрипт уже разбирает сохранённые файлы.

Использование:
    python3 tools/iafd_titles_to_docx.py страница.html [ещё.html ...] -o titles.docx
    python3 tools/iafd_titles_to_docx.py папка_с_html/ -o titles.docx
    python3 tools/iafd_titles_to_docx.py список.txt -o titles.docx   # по названию в строке

Опции:
    -o, --output   имя итогового .docx (по умолчанию iafd-titles.docx)
    --dedupe       убрать повторы названий внутри одного исполнителя
    --plain        только список названий, без года/студии/примечаний
"""

from __future__ import annotations

import argparse
import datetime as dt
import pathlib
import re
import sys

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

TITLE_HREF = re.compile(r"/title\.rme/", re.I)
YEAR = re.compile(r"^(19|20)\d{2}$")

# Заголовки вкладок фильмографии на странице IAFD -> человекочитаемая секция.
SECTION_NAMES = {
    "personal": "Фильмография",
    "peformer": "Как исполнитель",
    "performer": "Как исполнитель",
    "directoral": "Как режиссёр",
    "chapters": "Сцены / компиляции",
    "archive": "Архив",
}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def section_label(table) -> str:
    """Имя секции: по id таблицы, иначе по ближайшему заголовку выше."""
    tid = (table.get("id") or "").strip().lower()
    if tid in SECTION_NAMES:
        return SECTION_NAMES[tid]

    node = table
    for _ in range(6):
        node = node.find_previous(["h1", "h2", "h3", "h4", "a", "div"])
        if node is None:
            break
        label = clean(node.get_text())
        if 2 < len(label) < 60:
            return label
    return SECTION_NAMES.get(tid, "Фильмография")


def performer_name(soup: BeautifulSoup, fallback: str) -> str:
    for sel in ("h1", "title"):
        node = soup.find(sel)
        if node:
            name = clean(node.get_text())
            name = re.sub(r"\s*[-|]\s*IAFD.*$", "", name, flags=re.I)
            if name:
                return name
    return fallback


def parse_html(path: pathlib.Path) -> dict:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")

    sections: dict[str, list[dict]] = {}
    for row in soup.find_all("tr"):
        link = row.find("a", href=TITLE_HREF)
        if not link:
            continue

        cells = [clean(td.get_text(" ")) for td in row.find_all(["td", "th"])]
        title = clean(link.get_text(" ")) or (cells[0] if cells else "")
        if not title:
            continue

        rest = [c for c in cells if c and c != title]
        year = next((c for c in rest if YEAR.match(c)), "")
        rest = [c for c in rest if c != year]

        table = row.find_parent("table")
        label = section_label(table) if table else "Фильмография"
        sections.setdefault(label, []).append(
            {
                "title": title,
                "year": year,
                "distributor": rest[0] if rest else "",
                "notes": " / ".join(rest[1:]) if len(rest) > 1 else "",
                "url": link.get("href", ""),
            }
        )

    return {"name": performer_name(soup, path.stem), "source": path.name, "sections": sections}


def parse_txt(path: pathlib.Path) -> dict:
    rows = [
        {"title": line, "year": "", "distributor": "", "notes": "", "url": ""}
        for line in (clean(l) for l in path.read_text(encoding="utf-8", errors="replace").splitlines())
        if line
    ]
    return {"name": path.stem, "source": path.name, "sections": {"Фильмография": rows}}


def collect(inputs: list[str]) -> list[dict]:
    files: list[pathlib.Path] = []
    for raw in inputs:
        p = pathlib.Path(raw)
        if p.is_dir():
            files.extend(sorted(q for q in p.rglob("*") if q.suffix.lower() in {".html", ".htm", ".txt"}))
        elif p.is_file():
            files.append(p)
        else:
            print(f"пропускаю (не найдено): {p}", file=sys.stderr)

    docs = []
    for f in files:
        doc = parse_txt(f) if f.suffix.lower() == ".txt" else parse_html(f)
        total = sum(len(v) for v in doc["sections"].values())
        print(f"{f.name}: {doc['name']} — {total} назв.", file=sys.stderr)
        if total:
            docs.append(doc)
    return docs


def build_docx(docs: list[dict], out: pathlib.Path, dedupe: bool, plain: bool) -> int:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    document.add_heading("IAFD — списки названий", level=0)
    stamp = document.add_paragraph(
        f"Собрано {dt.datetime.now():%d.%m.%Y %H:%M} · исполнителей: {len(docs)}"
    )
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    grand_total = 0
    for doc in docs:
        document.add_page_break()
        document.add_heading(doc["name"], level=1)
        document.add_paragraph(f"источник: {doc['source']}")

        for label, rows in doc["sections"].items():
            if dedupe:
                seen, unique = set(), []
                for r in rows:
                    key = (r["title"].lower(), r["year"])
                    if key not in seen:
                        seen.add(key)
                        unique.append(r)
                rows = unique

            document.add_heading(f"{label} ({len(rows)})", level=2)
            grand_total += len(rows)

            if plain:
                for r in rows:
                    document.add_paragraph(r["title"], style="List Number")
                continue

            table = document.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            for cell, head in zip(table.rows[0].cells, ("Название", "Год", "Студия", "Примечания")):
                cell.paragraphs[0].add_run(head).bold = True
            for r in rows:
                cells = table.add_row().cells
                cells[0].text = r["title"]
                cells[1].text = r["year"]
                cells[2].text = r["distributor"]
                cells[3].text = r["notes"]

    document.save(out)
    return grand_total


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("inputs", nargs="+", help="сохранённые .html / .txt файлы или папка с ними")
    ap.add_argument("-o", "--output", default="iafd-titles.docx")
    ap.add_argument("--dedupe", action="store_true", help="убрать повторяющиеся названия")
    ap.add_argument("--plain", action="store_true", help="только названия, без таблицы")
    args = ap.parse_args()

    docs = collect(args.inputs)
    if not docs:
        print("Названий не найдено — проверь, что HTML сохранён целиком.", file=sys.stderr)
        return 1

    out = pathlib.Path(args.output)
    total = build_docx(docs, out, args.dedupe, args.plain)
    print(f"\nГотово: {out} — {total} названий из {len(docs)} страниц.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
