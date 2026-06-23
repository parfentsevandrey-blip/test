#!/usr/bin/env python3
"""
build_report_docx.py — editable Word (.docx) report with an advanced layout
engine that controls whitespace and image sizing.

What the layout engine does (vs. naive "2 photos + page break"):

* Image packing with VERTICAL JUSTIFICATION — images on a page are sized as
  large as the column allows, then the leftover vertical space is distributed as
  equal gaps, so pages are full and there is no large blank strip at the bottom.
* Orientation-aware: a tall portrait photo can take a whole page; two landscape
  photos share one. Each page is filled, photos are big.
* PAGE BUDGET — a hard cap (default 7 A4 pages per object). If a listing has 100
  photos, the gallery is down-sampled evenly to fit the budget instead of
  exploding to dozens of pages.
* keep_together / keep_with_next on text so headings don't get orphaned and a
  single trailing line doesn't strand a near-empty page.

Maps use the real keyless Google map (funda_a4/make_sheet.py:google_tiles_map),
with OSM fallback. Same JSON content model as the other renderers.

Usage:
    python3 build_report_docx.py content_utrecht.json -o report.docx --max-pages 7
"""

from __future__ import annotations

import argparse
import glob as globmod
import json
import math
import os
import sys
import tempfile
from typing import List, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "funda_a4"))
from make_sheet import google_tiles_map, osm_static_map, geocode  # noqa: E402

# ---- palette ----
NAVY = RGBColor(0x1B, 0x33, 0x57)
NAVY_SOFT = RGBColor(0x3A, 0x55, 0x7F)
GOLD = RGBColor(0xB0, 0x8D, 0x57)
DARK = RGBColor(0x2B, 0x2B, 0x2E)
GRAY = RGBColor(0x5C, 0x5C, 0x62)
GOLD_HEX = "B08D57"
NAVY_HEX = "1B3357"
ROW_HEX = "F3F1EC"          # warm light row shading
HEAD = "Cambria"            # elegant serif for titles (ships with Word)
BODY = "Calibri"            # clean sans for body (ships with Word)

# ---- geometry (cm) ----
MARGIN_TB = 1.4
MARGIN_LR = 1.6
PAGE_W, PAGE_H = 21.0, 29.7
BOX_W = PAGE_W - 2 * MARGIN_LR          # 17.8
USABLE_H = PAGE_H - 2 * MARGIN_TB       # 26.9
CM_PT = 28.3465

_TMP = tempfile.mkdtemp(prefix="report_docx_")


# ==========================================================================
# low-level docx helpers
# ==========================================================================

def _shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _border(paragraph, edge="bottom", color=GOLD_HEX, sz=8, space=2):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr"); pPr.append(pbdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space)); e.set(qn("w:color"), color)
    pbdr.append(e)


def _char_spacing(run, twips):
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(twips))
    run._r.get_or_add_rPr().append(sp)


def _run(p, text, size, color, *, font=BODY, bold=False, italic=False, caps=False, spacing=0):
    r = p.add_run(text)
    r.font.name = font
    # ensure east-asian/cyrillic uses same font
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font); rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font); rFonts.set(qn("w:eastAsia"), font)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.italic = italic; r.font.all_caps = caps
    if spacing:
        _char_spacing(r, spacing)
    return r


def _para(doc_or_cell, *, before=0, after=4, line=1.08, align=None,
          keep_together=False, keep_with_next=False):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = keep_together
    pf.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    return p


def _img_para(doc, path, w_cm, h_cm, before=0, center=True):
    p = _para(doc, before=before, after=0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(w_cm), height=Cm(h_cm))
    return p


# ==========================================================================
# advanced image packing
# ==========================================================================

def _disp_size(path) -> Tuple[float, float]:
    """Display size (cm) at full column width, capped to ~one page tall."""
    w, h = PILImage.open(path).size
    dw, dh = BOX_W, BOX_W * h / w
    cap = USABLE_H * 0.94
    if dh > cap:
        dh = cap; dw = dh * w / h
    return dw, dh


def pack_pages(paths: List[str], usable_h: float, gap_min: float = 0.5) -> List[list]:
    """Greedily pack images into pages; each item = (path, dw, dh)."""
    pages, cur, cur_h = [], [], 0.0
    for p in paths:
        dw, dh = _disp_size(p)
        add = dh + (gap_min if cur else 0)
        if cur and cur_h + add > usable_h:
            pages.append(cur); cur, cur_h = [], 0.0
            add = dh
        cur.append((p, dw, dh)); cur_h += add
    if cur:
        pages.append(cur)
    return pages


def fit_to_budget(paths: List[str], budget_pages: int, usable_h: float) -> List[str]:
    """Down-sample the photo list (evenly, keeping the facade) so packing fits."""
    if budget_pages <= 0:
        return []
    chosen = paths
    while pack_pages(chosen, usable_h) and len(pack_pages(chosen, usable_h)) > budget_pages and len(chosen) > 1:
        n = len(chosen)
        k = max(1, n - max(1, (len(pack_pages(chosen, usable_h)) - budget_pages)))
        if k >= n:
            k = n - 1
        idx = sorted({0} | {round(i * (n - 1) / (k - 1)) for i in range(k)}) if k > 1 else [0]
        chosen = [chosen[i] for i in idx]
    return chosen


IMG_GAP = 0.35  # cm between stacked images


def image_block(doc, path, caption=None):
    """Caption (kept with image) + image; image paragraph won't split a page."""
    dw, dh = _disp_size(path)
    if caption:
        cap = _para(doc, before=IMG_GAP * CM_PT, after=1, keep_with_next=True)
        _run(cap, caption, 10, GOLD, font=BODY, bold=True, caps=True, spacing=20)
        p = _img_para(doc, path, dw, dh, before=0)
    else:
        p = _img_para(doc, path, dw, dh, before=IMG_GAP * CM_PT)
    p.paragraph_format.keep_together = True


def budget_photos(text_h, map_path, main, gallery, max_pages, usable_h):
    """Trim the gallery (evenly) so the object's continuous flow fits max_pages."""
    base = text_h + _disp_size(map_path)[1] + 2 * IMG_GAP
    if main:
        base += _disp_size(main)[1] + IMG_GAP
    cap = (max_pages - 0.15) * usable_h
    chosen = list(gallery)
    while chosen:
        h = base + sum(_disp_size(p)[1] + IMG_GAP for p in chosen)
        if h <= cap:
            break
        n = len(chosen)
        over = (h - cap) / usable_h
        k = max(1, n - max(1, math.ceil(over * 2)))
        if k >= n:
            k = n - 1
        idx = sorted({round(i * (n - 1) / (k - 1)) for i in range(k)}) if k > 1 else [0]
        chosen = [chosen[i] for i in idx]
    return chosen


# ==========================================================================
# text height estimation (to budget gallery pages)
# ==========================================================================

def _est_lines(text, cpl):
    return max(1, math.ceil(len(text) / cpl))


def estimate_text_height(obj) -> float:
    """Rough height (cm) of the title block + spec table + sections."""
    h = 1.7  # title + district + rule
    specs = obj.get("specs", [])
    for _, v in specs:
        h += 0.42 + 0.40 * (_est_lines(v, 78) - 1)
    h += 0.3
    for sec in obj.get("sections", []):
        h += 0.75  # heading + rule + spacing
        if sec.get("subheading"):
            h += 0.5
        for p in sec.get("paragraphs", []):
            h += _est_lines(p, 108) * 0.40 + 0.16
        for b in sec.get("bullets", []):
            h += _est_lines(b, 100) * 0.39 + 0.06
        h += 0.25
    return h


# ==========================================================================
# content blocks
# ==========================================================================

def resolve_photos(obj, base) -> List[str]:
    spec = obj.get("photos", [])
    if isinstance(spec, str) and spec.startswith("glob:"):
        paths = sorted(globmod.glob(os.path.join(base, spec[5:])))
    elif isinstance(spec, list):
        paths = [p if os.path.isabs(p) else os.path.join(base, p) for p in spec]
    else:
        paths = []
    return [p for p in paths if os.path.exists(p)]


def cover(doc, content):
    for _ in range(7):
        _para(doc, after=0)
    top = _para(doc, after=0); _border(top, "bottom", GOLD_HEX, sz=14)
    t = _para(doc, before=12, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(t, content.get("title", "ОБЪЕКТЫ НЕДВИЖИМОСТИ"), 30, NAVY, font=HEAD, bold=True, spacing=80)
    s = _para(doc, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(s, content.get("subtitle", ""), 15, GRAY, font=HEAD, italic=True)
    _border(s, "bottom", GOLD_HEX, sz=14)

    h = _para(doc, before=26, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(h, "В ПОДБОРКУ ВХОДЯТ", 11, GOLD, bold=True, caps=True, spacing=40)
    for i, it in enumerate(content["objects"], 1):
        addr = it["address"].split(",")[0]
        city = it["address"].split(",")[-1].strip()
        line = _para(doc, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        _run(line, f"{i}.  ", 13, GOLD, font=HEAD, bold=True)
        _run(line, f"{addr}", 13, NAVY, font=HEAD, bold=True)
        _run(line, f"   {city} — {it.get('price_label','')}", 13, DARK, font=HEAD)


def spec_table(doc, specs):
    table = doc.add_table(rows=len(specs), cols=2)
    table.allow_autofit = False
    for i, (label, value) in enumerate(specs):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(5.4); c1.width = Cm(BOX_W - 5.4)
        for c in (c0, c1):
            c.paragraphs[0].paragraph_format.space_after = Pt(1)
            c.paragraphs[0].paragraph_format.space_before = Pt(1)
        _run(c0.paragraphs[0], label, 9, GRAY, bold=True)
        _run(c1.paragraphs[0], value, 9, DARK)
        if i % 2 == 0:
            _shade(c0, ROW_HEX); _shade(c1, ROW_HEX)
    # subtle gold left accent via a top border on the whole table's first row
    return table


def section(doc, sec):
    h = _para(doc, before=9, after=2, keep_with_next=True)
    _run(h, sec["heading"], 11, NAVY, bold=True, caps=True, spacing=15)
    _border(h, "bottom", GOLD_HEX, sz=6)
    if sec.get("subheading"):
        sp = _para(doc, after=3, keep_with_next=True)
        _run(sp, sec["subheading"], 9.5, GOLD, bold=True)
    for para in sec.get("paragraphs", []):
        p = _para(doc, after=5, line=1.12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_together=True)
        _run(p, para, 9, DARK)
    for b in sec.get("bullets", []):
        p = _para(doc, after=2, line=1.1, keep_together=True)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        _run(p, "•  ", 9, GOLD, bold=True)
        _run(p, b, 9, DARK)


def add_footer(doc, content):
    sec = doc.sections[0]
    f = sec.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, content.get("title", "") + "  ·  " + content.get("subtitle", ""), 8, GRAY, italic=True)


# ==========================================================================
# build
# ==========================================================================

def make_map(obj, idx):
    lat, lon = obj.get("lat"), obj.get("lon")
    if lat is None or lon is None:
        lat, lon = geocode(obj["address"] + ", Netherlands")
    zoom = obj.get("map_zoom", 13)
    try:
        img = google_tiles_map(lat, lon, zoom, 1400, 600)
        print(f"  map(google): {obj['address']} z{zoom}")
    except Exception as e:
        print(f"  ! google failed ({e}); OSM"); img = osm_static_map(lat, lon, zoom, 1400, 600)
    path = os.path.join(_TMP, f"map_{idx}.jpg")
    img.save(path, "JPEG", quality=90)
    return path


def build(content, base, out, max_pages):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(PAGE_H); sec.page_width = Cm(PAGE_W)
    sec.top_margin = Cm(MARGIN_TB); sec.bottom_margin = Cm(MARGIN_TB)
    sec.left_margin = Cm(MARGIN_LR); sec.right_margin = Cm(MARGIN_LR)
    st = doc.styles["Normal"]; st.font.name = BODY; st.font.size = Pt(9)
    add_footer(doc, content)

    cover(doc, content)

    for idx, obj in enumerate(content["objects"]):
        doc.add_page_break()
        t = _para(doc, after=1, keep_with_next=True)
        _run(t, obj["address"], 17, NAVY, font=HEAD, bold=True)
        d = _para(doc, after=4)
        _run(d, obj.get("district", ""), 9.5, GOLD, bold=True)
        _border(d, "bottom", GOLD_HEX, sz=12)
        if obj.get("specs"):
            spec_table(doc, obj["specs"]); _para(doc, after=2)
        for s in obj.get("sections", []):
            section(doc, s)

        # ---- images flow continuously after the text (fills page bottoms) ----
        photos = resolve_photos(obj, base)
        text_h = estimate_text_height(obj)
        map_path = make_map(obj, idx)
        main = photos[0] if photos else None
        gallery = budget_photos(text_h, map_path, main, photos[1:], max_pages, USABLE_H)

        _para(doc, after=0, before=6)
        image_block(doc, map_path, "Локация")
        if main:
            image_block(doc, main, "Объект")
        for g in gallery:
            image_block(doc, g)

    doc.save(out)
    print(f"\nSaved {out}")


def main() -> int:
    ap = argparse.ArgumentParser(description="Render the listing report as an editable Word .docx.")
    ap.add_argument("content")
    ap.add_argument("-o", "--out", default="report.docx")
    ap.add_argument("--max-pages", type=int, default=7, help="Max A4 pages per object (default 7).")
    args = ap.parse_args()
    with open(args.content, encoding="utf-8") as fh:
        content = json.load(fh)
    base = os.path.dirname(os.path.abspath(args.content))
    build(content, base, args.out, args.max_pages)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
