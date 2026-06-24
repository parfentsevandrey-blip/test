#!/usr/bin/env python3
"""
fill_registry.py — enrich a commercial-registry .docx:

1) Object photo pages: after each "Вставьте фотографии объекта ниже" anchor,
   insert the Google map (map_i.jpg) + main Funda photo (photo_i.jpg); the
   instruction line is removed. Images only, no captions.

2) Zone pages (Part 3): after each "Объекты реестра в этой зоне" anchor, put the
   zone's media on its OWN page (clean page_break_before — no blank pages):
   the Google location map (zmap_k.jpg) + a ground photo of the zone
   (zphoto_k.jpg) if available, otherwise an aerial (zaer_k.jpg).

Everything else in the document is left untouched.
"""
import sys, os, re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from PIL import Image

SRC, OUT, ASSETS = sys.argv[1], sys.argv[2], sys.argv[3]
MAX_W = 17.4
PH_ANCHOR = "Вставьте фотографии объекта ниже"
ZONE_ANCHOR = "Объекты реестра в этой зоне"


def insert_after(par):
    new_p = OxmlElement("w:p")
    par._p.addnext(new_p)
    return Paragraph(new_p, par._parent)


def set_image(p, path, max_w, max_h, space_after=6, page_break_before=False):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_together = True
    p.paragraph_format.page_break_before = page_break_before
    w, h = Image.open(path).size
    tw = max_w; th = tw * h / w
    if th > max_h:
        th = max_h; tw = th * w / h
    p.add_run().add_picture(path, width=Cm(tw), height=Cm(th))
    return p


def main():
    doc = Document(SRC)
    paras = doc.paragraphs

    # ---- 1) object photo pages ----
    ph_anchors = [p for p in paras if p.text.strip().startswith(PH_ANCHOR)]
    for i, anchor in enumerate(ph_anchors, 1):
        cur = anchor
        mp = os.path.join(ASSETS, f"map_{i}.jpg")
        ph = os.path.join(ASSETS, f"photo_{i}.jpg")
        if os.path.exists(mp):
            cur = set_image(insert_after(cur), mp, MAX_W, 9.6)
        if os.path.exists(ph):
            cur = set_image(insert_after(cur), ph, MAX_W, 10.2)
        anchor._p.getparent().remove(anchor._p)
        print(f"obj{i}: map={'+' if os.path.exists(mp) else '-'} photo={'+' if os.path.exists(ph) else '-'}")

    # ---- 2) zone pages ----
    # zone headings ("N. Name" followed by a "City, провинция/регион ..." line)
    headings = []
    for idx in range(len(paras) - 1):
        t = paras[idx].text.strip()
        nxt = paras[idx + 1].text
        if re.match(r"^\d+\.\s+\S", t) and len(t) < 80 and ("провинция" in nxt or "регион" in nxt):
            headings.append(paras[idx])

    zone_anchors = [p for p in paras if p.text.strip().startswith(ZONE_ANCHOR)]
    for k, anchor in enumerate(zone_anchors, 1):
        zmap = os.path.join(ASSETS, f"zmap_{k}.jpg")
        zphoto = os.path.join(ASSETS, f"zphoto_{k}.jpg")
        zaer = os.path.join(ASSETS, f"zaer_{k}.jpg")
        zpic = zphoto if os.path.exists(zphoto) else (zaer if os.path.exists(zaer) else None)
        cur = anchor
        if os.path.exists(zmap):
            cur = set_image(insert_after(cur), zmap, MAX_W, 9.6, page_break_before=True)  # media on new page
        if zpic:
            cur = set_image(insert_after(cur), zpic, MAX_W, 10.8)
        # next zone starts on a fresh page (no empty paragraph -> no blank page)
        if k < len(headings):
            headings[k].paragraph_format.page_break_before = True
        kind = "ground" if os.path.exists(zphoto) else ("aerial" if os.path.exists(zaer) else "none")
        print(f"zone{k}: map={'+' if os.path.exists(zmap) else '-'} photo={kind}")

    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()
